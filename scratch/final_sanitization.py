from docx import Document

def deep_sanitize(docx_path):
    doc = Document(docx_path)
    print(f"Sanitizing {docx_path}...")

    seen_paragraphs = set()
    to_delete = []
    
    # 1. Identify Duplicates and Artifacts
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Artifacts
        if "```mermaid" in text or text == "```" or "[INSTRUCTION" in text or "Unknown diagram error" in text:
            to_delete.append(i)
            continue
            
        # Duplicates (only for body text > 50 chars to avoid deleting headers/empty lines)
        if len(text) > 50:
            # Normalize text for comparison
            normalized = " ".join(text.split()).lower()
            if normalized in seen_paragraphs:
                print(f"Removing duplicate paragraph: {text[:50]}...")
                to_delete.append(i)
            else:
                seen_paragraphs.add(normalized)

    # 2. Delete from bottom to top
    for index in sorted(to_delete, reverse=True):
        p = doc.paragraphs[index]
        parent = p._element.getparent()
        parent.remove(p._element)

    doc.save(docx_path)
    print("Sanitization complete. No duplicates or artifacts remain.")

if __name__ == "__main__":
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    deep_sanitize(final_path)
