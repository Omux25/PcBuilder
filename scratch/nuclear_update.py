import os
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def nuclear_update():
    downloads = r'c:\Users\Omux2\Downloads'
    original_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport.docx'
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    
    shutil.copy2(original_path, final_path)
    doc = Document(final_path)
    print("Starting Nuclear Update...")

    # Aggressive Asset Mapping
    facade_img = os.path.join(downloads, "EMSI-Maroc-0x280.jpg")
    organigram_img = os.path.join(puml_dir, "emsi_organigramme.png")
    
    uml_map = {
        "Cas d'utilisation": os.path.join(puml_dir, "use_case.png"),
        "Diagramme de Classes": os.path.join(puml_dir, "class.png"),
        "moteur de compatibilité": os.path.join(puml_dir, "sequence_compatibility.png"),
        "collecte automatique": os.path.join(puml_dir, "sequence_scraping.png"),
        "Base de Données": os.path.join(assets_dir, "figure_6_database_erd_1778640799727.png")
    }

    # 1. Clean Mermaid and Broken Artifacts first
    in_mermaid = False
    to_delete = []
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t.startswith("```mermaid") or t == "```":
            in_mermaid = not in_mermaid
            to_delete.append(i)
            continue
        if in_mermaid or "Unknown diagram error" in t or "[INSTRUCTION" in t:
            to_delete.append(i)
            continue

    for index in sorted(to_delete, reverse=True):
        p = doc.paragraphs[index]
        parent = p._element.getparent()
        parent.remove(p._element)

    # 2. Insert EMSI Assets by keyword
    for para in doc.paragraphs:
        t = para.text.lower()
        # Facade
        if "façade" in t or "campus" in t or "orangers" in t:
            if os.path.exists(facade_img):
                para.text = ""
                run = para.add_run()
                run.add_picture(facade_img, width=Inches(5.0))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print("Inserted Facade.")
        # Organigram
        if "organigramme" in t:
            if os.path.exists(organigram_img):
                para.text = ""
                run = para.add_run()
                run.add_picture(organigram_img, width=Inches(5.0))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print("Inserted Organigram.")

    # 3. Insert UML Diagrams by keyword
    for para in doc.paragraphs:
        for key, path in uml_map.items():
            if key in para.text and os.path.exists(path):
                # Ensure we don't insert multiple times if keywords repeat
                # We'll just look for placeholders or section starts
                if "Modélisation" in para.text or "Figure" in para.text:
                    run = para.add_run()
                    run.add_break()
                    run.add_picture(path, width=Inches(5.2))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break

    # 4. Final Formatting
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    doc.save(final_path)
    print("Nuclear update complete.")

if __name__ == "__main__":
    nuclear_update()
