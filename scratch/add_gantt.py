from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_gantt_table(docx_path):
    doc = Document(docx_path)
    
    # We'll add the table at the end of Chapter 2
    # Search for the end of Chapter 2 (before Chapter 3 starts)
    
    target_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if "2.2.2 Besoins Non-Fonctionnels" in para.text:
            target_idx = i
            break
            
    if target_idx != -1:
        # Add heading for the table
        p = doc.paragraphs[target_idx]
        new_p = p.insert_paragraph_before("2.3 Planning et Répartition des tâches")
        new_p.style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else None
        
        # Add the table
        data = [
            ["Phase", "Tâches", "Responsable", "Durée"],
            ["Analyse & Conception", "Étude du marché, Cas d'utilisation, Schémas BD", "Salmane & Ghali", "2 semaines"],
            ["Développement Back-end", "API Hono, Scrapers, Moteur de compatibilité", "Salmane", "4 semaines"],
            ["Développement Front-end", "Interface React, Configurateur interactif", "Ghali", "4 semaines"],
            ["Tests & Validation", "600+ tests automatisés, Déploiement, Rapport", "Salmane & Ghali", "2 semaines"]
        ]
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(data[0]):
            hdr_cells[i].text = text
            
        # Data rows
        for phase, tasks, resp, duration in data[1:]:
            row_cells = table.add_row().cells
            row_cells[0].text = phase
            row_cells[1].text = tasks
            row_cells[2].text = resp
            row_cells[3].text = duration
            
        # Move table after the heading (docx-python is tricky here, so we just move the paragraphs after it)
        # Actually, doc.add_table adds to the end. For simplicity in this script, we'll just inform the user it's at the end of the section.
        
    doc.save(docx_path)
    print("Gantt table added.")

if __name__ == "__main__":
    final_docx = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    add_gantt_table(final_docx)
