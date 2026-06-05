import os
import shutil
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def definitive_update(docx_path, assets_map):
    # Start fresh from original
    original_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport.docx'
    final_path = docx_path
    shutil.copy2(original_path, final_path)
    
    doc = Document(final_path)
    print(f"Applying definitive updates to {final_path}")
    
    # Text expansion content
    ch2_text = """
2.1 Problématique
Le marché du matériel informatique au Maroc se caractérise par une asymétrie d'information majeure. Un utilisateur souhaitant assembler un ordinateur doit faire face à deux obstacles critiques :
1. La fragmentation de l'offre : Les revendeurs (UltraPC, NextLevel, SetupGame, etc.) opèrent de manière isolée, rendant la comparaison manuelle des prix extrêmement chronophage.
2. La complexité technique : L'interdépendance des composants (compatibilité du socket CPU, puissance de l'alimentation, dimensions du boîtier) crée un risque élevé d'erreurs d'achat coûteuses pour les néophytes.

2.2 Cahier des Charges
Le projet vise à fournir une solution intégrée de configuration et de comparaison.

2.2.1 Besoins Fonctionnels
- Pour l'Utilisateur :
    - Recherche et filtrage multicritères par catégorie de composants.
    - Configurateur assisté avec vérification en temps réel de la compatibilité.
    - Comparaison dynamique des prix chez les principaux revendeurs marocains.
    - Visualisation de l'historique des prix pour optimiser le moment de l'achat.
- Pour l'Administrateur :
    - Gestion du catalogue centralisé (CRUD des composants).
    - Supervision des sessions de collecte automatique (Scraping).
    - Mapping manuel des produits non reconnus automatiquement par le système.
    - Consultation des journaux d'erreurs techniques.

2.2.2 Besoins Non-Fonctionnels
- Performance : Le temps de réponse pour la validation d'une configuration doit être inférieur à 200ms pour garantir une expérience fluide.
- Sécurité : L'accès à l'interface d'administration doit être protégé par une authentification JWT (JSON Web Token) avec hachage Bcrypt des mots de passe.
- Fiabilité : Le système doit être capable de gérer les changements de structure HTML des sites sources grâce à des sélecteurs CSS robustes.
- Scalabilité : L'architecture doit permettre l'ajout de nouveaux revendeurs sans modification du moteur de compatibilité.
"""

    ch3_text = """
3.1 Justification des Choix Techniques
Le choix de la stack technologique a été dicté par des impératifs de performance et de modernité.

- React & TypeScript : L'utilisation de React permet une gestion efficace de l'état du configurateur. TypeScript apporte un typage statique crucial pour sécuriser les manipulations de spécifications techniques complexes (TDP, Sockets, Form Factors).
- Bun & Hono : Bun a été choisi comme "runtime" pour sa rapidité supérieure à Node.js, particulièrement lors de l'exécution des tests automatisés. Hono offre une API légère et extrêmement performante, idéale pour les services de validation en temps réel.
- PostgreSQL : Pour la persistence des données, PostgreSQL a été retenu pour sa gestion robuste des relations et son support natif du format JSONB, nécessaire pour stocker les fiches techniques variées des composants.

3.2 Modélisation de l'Architecture
L'application suit une architecture modulaire composée de trois piliers :
1. Le moteur de collecte (Scrapers) : Découple la récupération des données de leur traitement.
2. L'agrégateur & DNA Matcher : Identifie et lie les offres brutes aux composants du catalogue via un algorithme de correspondance textuelle.
3. Le moteur de compatibilité : Un service purement logique qui valide les règles métier indépendamment de l'interface.
"""

    paragraphs = doc.paragraphs
    to_delete = []
    in_mermaid = False
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        
        # 1. Clean Mermaid
        if text.startswith("```mermaid") or text == "```":
            in_mermaid = not in_mermaid
            to_delete.append(i)
            continue
        if in_mermaid:
            to_delete.append(i)
            continue
            
        # 2. Expansion Logic
        if "2.1 Problématique" in para.text and len(para.text) < 150:
            para.text = ch2_text
        if "3.1 Analyse des besoins" in para.text and len(para.text) < 150:
            para.text = ch3_text
            
        # 3. Asset Insertion
        for placeholder, img_path in assets_map.items():
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, "")
                if os.path.exists(img_path):
                    run = para.add_run()
                    run.add_break()
                    run.add_picture(img_path, width=Inches(5.0))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

    # Delete marked paragraphs
    for index in sorted(to_delete, reverse=True):
        p = doc.paragraphs[index]
        parent = p._element.getparent()
        parent.remove(p._element)

    # 4. Add Gantt at the end of Chapter 2
    # (Finding the new insertion point after expansion)
    for i, para in enumerate(doc.paragraphs):
        if "2.2.2 Besoins Non-Fonctionnels" in para.text:
            new_p = para.insert_paragraph_before("2.3 Planning et Répartition des tâches")
            new_p.style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else None
            
            table = doc.add_table(rows=5, cols=4)
            table.style = 'Table Grid'
            data = [
                ["Phase", "Tâches", "Responsable", "Durée"],
                ["Analyse & Conception", "Étude du marché, Cas d'utilisation, Schémas BD", "Salmane & Ghali", "2 semaines"],
                ["Développement Back-end", "API Hono, Scrapers, Moteur de compatibilité", "Salmane", "4 semaines"],
                ["Développement Front-end", "Interface React, Configurateur interactif", "Ghali", "4 semaines"],
                ["Tests & Validation", "600+ tests automatisés, Déploiement, Rapport", "Salmane & Ghali", "2 semaines"]
            ]
            for row_idx, row_data in enumerate(data):
                for col_idx, cell_text in enumerate(row_data):
                    table.cell(row_idx, col_idx).text = cell_text
            break

    doc.save(final_path)
    print("Definitive update complete.")

if __name__ == "__main__":
    downloads = r'c:\Users\Omux2\Downloads'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    
    assets_map = {
        "Insérer ici une image de la façade du campus EMSI Orangers": os.path.join(downloads, "EMSI-Maroc-0x280.jpg"),
        "Insérer ici l'Organigramme officiel": os.path.join(puml_dir, "emsi_organigramme.png"),
        "3.2 Modélisation UML : Cas d'utilisation": os.path.join(puml_dir, "use_case.png"),
        "3.3 Architecture et Diagramme de Classes": os.path.join(puml_dir, "class.png"),
        "3.4 Conception de la Base de Données": os.path.join(assets_dir, "figure_6_database_erd_1778640799727.png"),
        "Figure 7 : Lecture visuelle du moteur de compatibilité": os.path.join(puml_dir, "sequence_compatibility.png"),
        "Figure 8 : Chaîne de collecte automatique des prix": os.path.join(puml_dir, "sequence_scraping.png"),
        "Figure 9 : Aperçu fonctionnel des interfaces": os.path.join(assets_dir, "figure_7_home_1778642725154.png"),
        "Figure 10 : Synthèse visuelle des résultats": os.path.join(assets_dir, "figure_10_admin_dashboard_1778643047831.png")
    }
    
    definitive_update(r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx', assets_map)
