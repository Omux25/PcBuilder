import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_ultimate_technical_report():
    downloads = r'c:\Users\Omux2\Downloads'
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final_V2.docx'
    
    with open(r'c:\Headquarters\Projects\PcBuilder\scratch\final_text_audit.txt', 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    print("Building THE ULTIMATE TECHNICAL EDITION (50+ Pages)...")

    # --- HELPERS ---
    def add_centered(text, size=12, bold=False, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size); run.bold = bold
        p.paragraph_format.space_after = Pt(space_after)

    # --- TITLE PAGE ---
    add_centered("RAPPORT DE PROJET DE FIN D'ANNÉE", 24, True, 20)
    add_centered("PC Builder Maroc", 22, True, 10)
    add_centered("Plateforme de comparaison de prix et de vérification de compatibilité", 14, False, 40)
    add_centered("École : EMSI Orangers, Casablanca", 12, False, 5)
    add_centered("Filière : 3IIR — Ingénierie Informatique et Réseaux", 12, False, 20)
    add_centered("Étudiants : Salmane ELHJOUJI · Ghali KHARMOUDY", 12, False, 5)
    add_centered("Encadrante : Prof. Houda Mouttalib", 12, False, 20)
    add_centered("Année : 2025/2026", 12)
    doc.add_page_break()

    # --- TECHNICAL CONTENT BLOCKS ---
    
    comp_intro = "Le moteur de compatibilité est le cœur décisionnel de PC Builder Maroc. Contrairement à une simple base de données, il implémente un ensemble de règles métier complexes qui traduisent les contraintes physiques du matériel informatique en logique booléenne."
    
    rules_analysis = [
        ("Règle 1 : Incompatibilité de Socket", "Cette règle critique compare le champ 'socket' du CPU et de la carte mère. Un processeur LGA1700 ne peut physiquement pas être installé sur une carte AM5. Le système bloque immédiatement la configuration."),
        ("Règle 2 : Type de Mémoire Vive (DDR4 vs DDR5)", "Les slots DIMM sont physiquement différents. Le moteur vérifie que le champ 'supported_ram_types' de la carte mère inclut le type de la RAM sélectionnée."),
        ("Règle 3 : Fréquence Mémoire Maximale", "Si la RAM dépasse la fréquence maximale supportée par la carte mère, le système génère un avertissement technique indiquant que la mémoire sera bridée à la vitesse du bus système."),
        ("Règle 4 : Dimensionnement de la Carte Graphique", "Le système compare la longueur (mm) du GPU au dégagement maximal (max_gpu_length_mm) du boîtier. C'est une erreur physique bloquante."),
        ("Règle 5 : Facteur de Forme (Form Factor)", "Vérification de la compatibilité entre le format de la carte mère (ATX, Micro-ATX, ITX) et les formats supportés par le boîtier."),
        ("Règle 6 : Hauteur du Ventirad", "Comparaison de la hauteur du système de refroidissement avec la largeur maximale interne du boîtier."),
        ("Règle 7 : Disponibilité des Slots DIMM", "Le système compte le nombre de barrettes (y compris dans les kits) et le compare aux slots physiques disponibles sur la carte mère."),
        ("Règle 8 : Slots de Stockage (M.2 & SATA)", "Analyse combinée du nombre de disques NVMe et SATA par rapport aux ports disponibles sur le chipset."),
        ("Règle 9 : Support Socket du Refroidisseur", "Vérification que les fixations du ventirad ou de l'AIO sont compatibles avec le socket du CPU choisi."),
        ("Règle 10 : Mixage des Types de RAM", "Interdiction stricte de mélanger différents types de RAM dans une même configuration pour garantir la stabilité du système."),
        ("Règle 11 : Mixage des Fréquences", "Avertissement sur le fait que le système s'alignera sur la fréquence la plus basse du pool de mémoire."),
        ("Règle 12 : Dissipation Thermique (TDP)", "Avertissement si le TDP du processeur dépasse la capacité de dissipation (Max TDP) du refroidisseur."),
        ("Règle 13 : Configuration Dual-Channel", "Optimisation suggérant d'utiliser des paires de barrettes pour doubler la bande passante mémoire."),
        ("Règle 14 : Facteur de Forme de l'Alimentation", "Vérification de la compatibilité entre le format du bloc d'alimentation (ATX, SFX) et le boîtier."),
        ("Règle 15 : Calcul de Charge Électrique (TDP Total)", "Calcul de la consommation cumulée avec une marge de sécurité de 50% (Coefficient 1.5) pour recommander l'alimentation idéale.")
    ]

    dna_matcher_analysis = """
L'algorithme DNA Matcher utilise une approche de tokenisation avancée pour normaliser les données brutes issues du web. Le processus se décompose en plusieurs étapes :
1. Décodage des entités HTML (gestion des caractères spéciaux des revendeurs).
2. Extraction de la marque via une liste d'autorité de 130+ constructeurs.
3. Suppression des 'Noise Tokens' (kit, bundle, oem, edition) qui polluent la correspondance.
4. Filtrage des 'Color Tokens' (noir, blanc, rgb) pour isoler le modèle technique pur.
5. Calcul de score par mot-clé pour classifier le produit dans l'une des 10 catégories matérielles.
"""

    # --- MAPPING ---
    sections = [
        ("Remerciements", lines[10:16]),
        ("Résumé", lines[17:21]),
        ("Résumé Technique (Abstract)", [
            "Ce projet concerne la conception et le développement de PC Builder Maroc, une plateforme web complète spécifiquement adaptée au marché marocain des composants informatiques. L'application permet aux utilisateurs de concevoir des configurations PC personnalisées tout en offrant une vérification automatique de la compatibilité technique et des comparaisons de prix en temps réel chez plusieurs revendeurs en ligne marocains.",
            "Construite sur une architecture Full-stack moderne, la solution utilise un frontend React, un serveur applicatif Bun/Hono et une base de données PostgreSQL. Elle dispose d'un moteur de compatibilité sophistiqué, d'un système de Scrapers automatiques avec suivi des tendances historiques et d'un Dashboard administratif sécurisé."
        ]),
        ("Table des matières", []),
        ("Introduction Générale", lines[32:36] + ["\nLe marché informatique marocain souffre d'une asymétrie d'information. Les revendeurs opèrent en silos, obligeant les clients à des comparaisons manuelles fastidieuses. Ce projet vise à centraliser ces flux de données."]),
        ("Chapitre 1 : Contexte du projet", []),
        ("1.1 Présentation de l’EMSI", lines[41:46]),
        ("1.2 Structure Organisationnelle", lines[48:51]),
        ("1.3 Analyse du Marché Marocain", lines[53:56] + ["\nL'émergence du e-commerce au Maroc, portée par des acteurs comme UltraPC et NextLevel, a créé un besoin de centralisation. Notre solution apporte cette couche d'intelligence métier indispensable."]),
        ("Chapitre 2 : Présentation du projet", []),
        ("2.1 Problématique", lines[63:67]),
        ("2.2 Cahier des Charges", lines[67:85]),
        ("2.3 Solution Proposée", lines[85:87]),
        ("2.4 Planning Détaillé", []),
        ("Chapitre 3 : Analyse et Conception", []),
        ("3.1 Choix Techniques et Stack", lines[91:96]),
        ("3.2 Architecture du Design System", ["\nL'interface utilise une architecture de composants atomiques. Chaque élément (bouton, card, modal) est standardisé pour garantir une cohérence visuelle sur le configurateur."]),
        ("3.3 Modélisation UML", lines[96:104]),
        ("Chapitre 4 : Réalisation technique", []),
        ("4.1 Environnement technique", lines[107:113]),
        ("4.2 Serveur applicatif (Back-end)", lines[113:119]),
        ("4.3 Le Moteur DNA Matcher", [dna_matcher_analysis]), # UNIQUE ANALYSIS
        ("4.4 Système de Scraping automatique", lines[128:135] + ["\nChaque scraper (UltraPC, NextLevel, SetupGame) possède une logique de parsing CSS dédiée, isolée dans des classes de services spécifiques."]),
        ("4.5 Moteur de Compatibilité : Analyse des 15 Règles", [comp_intro]), # MASSIVE EXPANSION START
    ]
    
    # Inject the 15 rules as separate sections
    for name, desc in rules_analysis:
        sections.append((name, [desc]))
        
    sections += [
        ("4.6 Dashboard d'administration", lines[145:151]),
        ("Chapitre 5 : Résultats, tests et validation", []),
        ("5.1 Méthodologie de Validation", lines[159:165] + ["\nNous utilisons des tests basés sur les propriétés avec 'fast-check' pour vérifier que le moteur de compatibilité reste stable face à des milliers de combinaisons de composants aléatoires."]),
        ("Conclusion et perspectives", lines[184:191]),
        ("Bibliographie", lines[192:206]),
        ("Annexes", lines[207:228])
    ]

    fig_count = 1
    def add_fig(path, title, width=5.0):
        nonlocal fig_count
        if os.path.exists(path):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(width))
            cap = doc.add_paragraph(f"Figure {fig_count} : {title}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_count += 1

    for title, content in sections:
        if "Chapitre" in title or title in ["Conclusion et perspectives", "Bibliographie", "Annexes"]:
            doc.add_page_break()
            doc.add_heading(title, level=1)
        elif title == "Table des matières":
            doc.add_heading(title, level=1)
            # Add placeholders
            for i in range(1, 40): doc.add_paragraph(f"Section {i} ................................................. {i}")
        else:
            doc.add_heading(title, level=2)
        
        for line in content:
            txt = line.strip()
            if txt: doc.add_paragraph(txt)
            
        if "Planning" in title: 
            table = doc.add_table(rows=5, cols=4); table.style = 'Table Grid'
            data = [["Phase", "Tâches", "Responsable", "Durée"], ["Analyse", "UML", "Binôme", "2 sem."], ["Back-end", "API", "Salmane", "4 sem."], ["Front-end", "UI", "Ghali", "4 sem."], ["Validation", "Tests", "Binôme", "2 sem."]]
            for r, row in enumerate(data):
                for c, val in enumerate(row): table.cell(r, c).text = val
            doc.add_paragraph(f"Figure {fig_count} : Planning GANTT").alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_count += 1
        if "Base de Données" in title: add_fig(os.path.join(puml_dir, "database_erd.png"), "Schéma Relationnel de la Base de Données")
        if "Analyse des 15 Règles" in title: add_fig(os.path.join(puml_dir, "sequence_compatibility.png"), "Flux de vérification de compatibilité")

    # Global Formatting
    for section in doc.sections:
        section.top_margin = Inches(0.98); section.bottom_margin = Inches(0.98)
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1.5
        if len(p.text) > 80: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            if not run.bold: run.font.size = Pt(12)

    doc.save(final_path)
    print("ULTIMATE TECHNICAL EDITION BUILT.")

if __name__ == "__main__":
    build_ultimate_technical_report()
