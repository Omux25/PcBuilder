import docx
from docx.shared import Pt

def check_formatting(file_path):
    doc = docx.Document(file_path)
    
    print(f"--- Formatting Audit for {file_path} ---")
    
    # Check default font/size (first paragraph as sample)
    if doc.paragraphs:
        p = doc.paragraphs[0]
        if p.runs:
            run = p.runs[0]
            print(f"Sample Font: {run.font.name}")
            print(f"Sample Size: {run.font.size.pt if run.font.size else 'Default'}")
    
    # Check margins
    section = doc.sections[0]
    print(f"Left Margin: {section.left_margin.cm:.2f} cm")
    print(f"Right Margin: {section.right_margin.cm:.2f} cm")
    print(f"Top Margin: {section.top_margin.cm:.2f} cm")
    print(f"Bottom Margin: {section.bottom_margin.cm:.2f} cm")
    
    # Check Line Spacing
    if doc.paragraphs:
        p = doc.paragraphs[10] # pick one with text
        print(f"Line Spacing: {p.paragraph_format.line_spacing}")

if __name__ == "__main__":
    path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport.docx'
    check_formatting(path)
