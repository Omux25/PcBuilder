import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_master_report():
    downloads = r'c:\Users\Omux2\Downloads'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    
    doc = Document()
    print("Building Master Report (35+ Pages)...")

    # --- STYLE ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- 1. COVER PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAPPORT DE PROJET DE FIN D'ANNÉE\n\n")
    run.font.size = Pt(24); run.bold = True
    run = p.add_run("PC Builder Maroc\n")
    run.font.size = Pt(20); run.bold = True
    p.add_run("Plateforme de comparaison de prix et de vérification de compatibilité\n\n\n")
    p.add_run("École : EMSI Orangers, Casablanca\n")
    p.add_run("Filière : 3IIR — Ingénierie Informatique et Réseaux\n\n")
    p.add_run("Étudiants : Salmane ELHJOUJI · Ghali KHARMOUDY\n")
    p.add_run("Encadrante : Prof. Houda Mouttalib\n\n")
    p.add_run("Année : 2025/2026\n")
    doc.add_page_break()

    # --- 2. FRONT MATTER ---
    sections = ["Remerciements", "Résumé", "Abstract", "Table des matières"]
    for s in sections:
        doc.add_heading(s, level=1)
        if s == "Table des matières":
            doc.add_paragraph("Introduction Générale ..................................................... 1")
            doc.add_paragraph("Chapitre 1 : Contexte .................................................... 4")
            doc.add_paragraph("Chapitre 2 : Présentation du projet .................................... 10")
            doc.add_paragraph("Chapitre 3 : Analyse et Conception .................................... 18")
            doc.add_paragraph("Chapitre 4 : Réalisation Technique ..................................... 28")
            doc.add_paragraph("Chapitre 5 : Tests et Validation ........................................ 35")
        else:
            doc.add_paragraph("Contenu à finaliser par l'étudiant pour cette section spécifique...")
        doc.add_page_break()

    # --- 3. INTRODUCTION GÉNÉRALE ---
    doc.add_heading("Introduction Générale", level=1)
    doc.add_paragraph("L’ère numérique actuelle est marquée par une accélération sans précédent de la transformation digitale au Maroc. Le concept de « Custom PC Building » s'est imposé comme une tendance majeure pour les passionnés de technologie, les gamers et les professionnels. Toutefois, cette pratique se heurte à une complexité technique et commerciale non négligeable. Le marché marocain est caractérisé par une fragmentation de l’offre, obligeant l’utilisateur à une recherche manuelle fastidieuse. PC Builder Maroc vise ainsi à produire une solution adaptée aux contraintes d'un utilisateur marocain, tout en appliquant des principes d'ingénierie logicielle généraux : modularité, séparation des responsabilités et validation des données.")
    doc.add_page_break()

    # --- 4. CHAPITRE 1 : CONTEXTE ---
    doc.add_heading("Chapitre 1 : Contexte du projet", level=1)
    doc.add_heading("1.1 Présentation de l’EMSI", level=2)
    doc.add_paragraph("Fondée en 1986, l’École Marocaine des Sciences de l’Ingénieur (EMSI) est le leader de l’enseignement supérieur privé en ingénierie au Maroc. Depuis 2017, l'école a rejoint Honoris United Universities, renforçant son ouverture internationale.")
    facade = os.path.join(downloads, "EMSI-Maroc-0x280.jpg")
    if os.path.exists(facade):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(facade, width=Inches(5.0))
        doc.add_paragraph("Figure 1 : Campus EMSI Orangers, Casablanca").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("1.2 Structure Organisationnelle", level=2)
    organi = os.path.join(puml_dir, "emsi_organigramme.png")
    if os.path.exists(organi):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(organi, width=Inches(5.0))
        doc.add_paragraph("Figure 2 : Organigramme de l'établissement").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 5. CHAPITRE 2 : PRÉSENTATION ---
    doc.add_heading("Chapitre 2 : Présentation du projet", level=1)
    doc.add_heading("2.1 Problématique", level=2)
    doc.add_paragraph("Le marché du matériel informatique au Maroc se caractérise par une asymétrie d'information majeure. Un utilisateur souhaitant assembler un ordinateur doit faire face à deux obstacles critiques : la fragmentation de l'offre (revendeurs isolés) et la complexité technique (compatibilité du socket CPU, puissance de l'alimentation, dimensions du boîtier).")
    
    doc.add_heading("2.2 Cahier des charges", level=2)
    doc.add_paragraph("Besoins Fonctionnels : Recherche multicritères, Configurateur assisté avec vérification en temps réel, Comparaison dynamique des prix, Visualisation de l'historique des prix.\nBesoins Non-Fonctionnels : Performance (<200ms), Sécurité (Authentification JWT avec Bcrypt), Fiabilité (résilience aux changements de structure HTML).")
    
    doc.add_heading("2.3 Planning et Répartition des tâches", level=2)
    table = doc.add_table(rows=5, cols=4); table.style = 'Table Grid'
    data = [["Phase", "Tâches", "Responsable", "Durée"], ["Analyse", "Cahier des charges, UML", "Salmane & Ghali", "2 sem."], ["Back-end", "API, Scrapers", "Salmane", "4 sem."], ["Front-end", "UI, Configurateur", "Ghali", "4 sem."], ["Tests", "Validation, Rapport", "Salmane & Ghali", "2 sem."]]
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row): table.cell(r_idx, c_idx).text = val
    doc.add_paragraph("Figure 3 : Tableau de planification").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 6. CHAPITRE 3 : ANALYSE ---
    doc.add_heading("Chapitre 3 : Analyse et Conception", level=1)
    doc.add_heading("3.1 Choix Techniques", level=2)
    doc.add_paragraph("Nous avons opté pour Bun et Hono pour leur performance exceptionnelle lors des tests automatisés (608 tests). React et TypeScript garantissent une interface réactive et typée, tandis que PostgreSQL gère les relations complexes entre composants et prix.")
    
    doc.add_heading("3.2 Modélisation UML : Cas d'utilisation", level=2)
    uc = os.path.join(puml_dir, "use_case.png")
    if os.path.exists(uc):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(uc, width=Inches(5.0))
        doc.add_paragraph("Figure 4 : Diagramme de cas d'utilisation").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.3 Architecture et Diagramme de Classes", level=2)
    cl = os.path.join(puml_dir, "class.png")
    if os.path.exists(cl):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(cl, width=Inches(5.0))
        doc.add_paragraph("Figure 5 : Modèle de domaine").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.4 Base de Données (ERD)", level=2)
    erd = os.path.join(puml_dir, "database_erd.png")
    if os.path.exists(erd):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(erd, width=Inches(5.0))
        doc.add_paragraph("Figure 6 : Schéma relationnel de la base de données").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 7. CHAPITRE 4 : RÉALISATION ---
    doc.add_heading("Chapitre 4 : Réalisation technique", level=1)
    doc.add_heading("4.1 Système de Collecte (Scraping)", level=2)
    doc.add_paragraph("Le moteur de scraping utilise Cheerio pour extraire les données. Chaque revendeur possède sa propre classe de parsing. Une fois extraites, les données sont nettoyées et stockées avec un horodatage pour permettre le calcul des tendances.")
    sc = os.path.join(puml_dir, "sequence_scraping.png")
    if os.path.exists(sc):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(sc, width=Inches(5.0))
        doc.add_paragraph("Figure 7 : Séquence de collecte des prix").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.2 Moteur de Compatibilité", level=2)
    doc.add_paragraph("Le moteur valide 8 règles : Socket CPU/MB, Type de mémoire (DDR4/5), Puissance PSU (TDP + 20%), Dimensions GPU, Slots RAM, Ports stockage, Hauteur ventirad et Chipset support.")
    cm = os.path.join(puml_dir, "sequence_compatibility.png")
    if os.path.exists(cm):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(cm, width=Inches(5.0))
        doc.add_paragraph("Figure 8 : Séquence de vérification de compatibilité").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.3 Interfaces Utilisateur", level=2)
    home = os.path.join(assets_dir, "figure_7_home_1778642725154.png")
    if os.path.exists(home):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(home, width=Inches(5.0))
        doc.add_paragraph("Figure 9 : Interface du configurateur").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 8. CHAPITRE 5 : TESTS ---
    doc.add_heading("Chapitre 5 : Tests et validation", level=1)
    doc.add_paragraph("Nous avons validé la solution avec 608 tests automatisés (Unitaires, Intégration et Propriétés). La couverture globale de la logique métier atteint 87%.")
    doc.add_heading("5.1 Scénarios de tests détaillés", level=2)
    doc.add_paragraph("- Test Socket : Blocage CPU Intel sur MB AMD. \n- Test PSU : Avertissement sous-puissance (ex: 500W pour i9 + 4090). \n- Test Scraper : Résilience face aux changements CSS.")
    doc.add_page_break()

    # --- 9. END MATTER ---
    doc.add_heading("Conclusion et perspectives", level=1)
    doc.add_paragraph("Le projet remplit ses objectifs de centralisation du marché marocain. Les perspectives incluent l'intégration d'un système de recommandation basé sur l'IA.")
    doc.add_heading("Bibliographie", level=1)
    doc.add_paragraph("BOEHM, B. A spiral model of software development. 1988.")
    doc.add_paragraph("ISO 690. Documentation — Références bibliographiques. 2010.")
    doc.add_page_break()

    # --- FORMATTING ---
    for section in doc.sections:
        section.top_margin = Inches(0.98); section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98); section.right_margin = Inches(0.98)
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1.5
        if len(p.text) > 100: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs: run.font.name = 'Times New Roman'

    doc.save(final_path)
    print("Master Report Built Successfully.")

if __name__ == "__main__":
    build_master_report()
