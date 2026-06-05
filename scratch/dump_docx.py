import docx
import sys

def dump_docx(path, out_path):
    doc = docx.Document(path)
    with open(out_path, "w", encoding="utf-8") as f:
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                f.write(f"[{i}] {p.style.name}: {text}\n")
        f.write("\n--- TABLES ---\n")
        for i, t in enumerate(doc.tables):
            f.write(f"Table {i}:\n")
            for row in t.rows:
                f.write(" | ".join([cell.text.strip() for cell in row.cells]) + "\n")

if __name__ == "__main__":
    dump_docx(sys.argv[1], sys.argv[2])
