import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_final_indexed_report():
    downloads = r'c:\Users\Omux2\Downloads'
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    
    with open(r'c:\Headquarters\Projects\PcBuilder\scratch\backup_content.txt', 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    print("Building FINAL INDEXED REPORT (No conflict numbers)...")

    # Title Page
    def add_centered(text, size=12, bold=False, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size); run.bold = bold
        p.paragraph_format.space_after = Pt(space_after)

    add_centered("RAPPORT DE PROJET DE FIN D'ANNÉE", 24, True, 20)
    add_centered("PC Builder Maroc", 22, True, 10)
    add_centered("Plateforme de comparaison de prix et de vérification de compatibilité", 14, False, 40)
    add_centered("École : EMSI Orangers, Casablanca", 12, False, 5)
    add_centered("Filière : 3IIR — Ingénierie Informatique et Réseaux", 12, False, 20)
    add_centered("Étudiants : Salmane ELHJOUJI · Ghali KHARMOUDY", 12, False, 5)
    add_centered("Encadrante : Prof. Houda Mouttalib", 12, False, 20)
    add_centered("Année : 2025/2026", 12)
    doc.add_page_break()

    # Cleaner Content Mapping
    sections = [
        ("Remerciements", lines[11:17]),
        ("Résumé", lines[18:22]),
        ("Abstract", lines[23:25]),
        ("Table des matières", []),
        ("Introduction Générale", lines[27:34] + lines[52:55]),
        ("Chapitre 1 : Contexte du projet", []),
        ("1.1 Présentation de l’EMSI", lines[35:38] + lines[39:41]),
        ("1.2 Structure Organisationnelle", lines[41:44]),
        ("1.3 Contexte du marché marocain", lines[44:47]),
        ("1.4 Problèmes observés", lines[47:52]),
        ("Chapitre 2 : Présentation du projet", []),
        ("2.1 Problématique", lines[63:68]),
        ("2.2 Cahier des Charges", lines[68:89]),
        ("2.3 Solution Proposée", lines[90:93]),
        ("2.4 Planning", []),
        ("Chapitre 3 : Analyse et Conception", []),
        ("3.1 Justification des Choix Techniques", lines[94:101]),
        ("3.2 Modélisation UML : Cas d'utilisation", []),
        ("3.3 Architecture et Diagramme de Classes", []),
        ("3.4 Conception de la Base de Données", []),
        ("Chapitre 4 : Réalisation technique", []),
        ("4.1 Environnement technique", lines[115:121]),
        ("4.2 Serveur applicatif et API REST", lines[121:127]),
        ("4.3 Moteur de compatibilité", lines[127:134]),
        ("4.4 Système de collecte automatique", lines[136:143]),
        ("4.5 Interface utilisateur", lines[145:152]),
        ("4.6 Panneau d'administration", lines[153:159]),
        ("Chapitre 5 : Résultats, tests et validation", []),
        ("5.1 Fonctionnalités livrées", lines[161:167]),
        ("5.2 Tests réalisés", lines[169:175]),
        ("5.3 Apport de la solution", lines[175:181]),
        ("5.4 Limites identifiées", lines[181:185]),
        ("5.5 Organisation du travail", lines[185:189]),
        ("5.6 Perspectives techniques détaillées", lines[189:194]),
        ("Conclusion et perspectives", lines[194:201]),
        ("Bibliographie", lines[202:216]),
        ("Annexes", lines[217:244])
    ]

    fig_count = 1
    def add_fig(path, title, width=5.0):
        nonlocal fig_count
        if os.path.exists(path):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(width))
            cap = doc.add_paragraph(f"Figure {fig_count} : {title}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_count += 1

    for title, content in sections:
        if "Chapitre" in title or title in ["Conclusion et perspectives", "Bibliographie", "Annexes"]:
            doc.add_page_break()
            doc.add_heading(title, level=1)
        elif title == "Table des matières":
            doc.add_heading(title, level=1)
            doc.add_paragraph("Introduction Générale ................................................. 1")
            doc.add_paragraph("Chapitre 1 : Contexte ................................................ 3")
            doc.add_paragraph("Chapitre 2 : Présentation du projet ................................ 9")
            doc.add_paragraph("Chapitre 3 : Analyse et Conception ................................ 16")
            doc.add_paragraph("Chapitre 4 : Réalisation Technique ................................. 24")
            doc.add_paragraph("Chapitre 5 : Tests et Validation .................................... 32")
        else:
            doc.add_heading(title, level=2)
        
        for line in content:
            txt = line.strip()
            # MANUALLY CLEAN STRAY HEADERS
            # Skip if it starts with "2.1", "2.2", etc. but NOT sub-headers like "2.2.1"
            if re.match(r'^\d+\.\d+\s', txt): continue
            # Skip if it matches the title
            if txt.lower() == title.lower(): continue
            if txt: doc.add_paragraph(txt)
            
        # Assets (Same logic as before but confirmed order)
        if "EMSI" in title: add_fig(os.path.join(downloads, "EMSI-Maroc-0x280.jpg"), "Campus EMSI Orangers")
        if "Organisationnelle" in title: add_fig(os.path.join(puml_dir, "emsi_organigramme.png"), "Organigramme")
        if "Planning" in title: 
            table = doc.add_table(rows=5, cols=4); table.style = 'Table Grid'
            data = [["Phase", "Tâches", "Responsable", "Durée"], ["Analyse", "UML", "Binôme", "2 sem."], ["Back-end", "API", "Salmane", "4 sem."], ["Front-end", "UI", "Ghali", "4 sem."], ["Validation", "Tests", "Binôme", "2 sem."]]
            for r, row in enumerate(data):
                for c, val in enumerate(row): table.cell(r, c).text = val
            doc.add_paragraph("Figure 3 : Planning").alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_count += 1
        if "Cas d'utilisation" in title: add_fig(os.path.join(puml_dir, "use_case.png"), "Diagramme de cas d'utilisation")
        if "Architecture" in title and "Classes" in title: add_fig(os.path.join(puml_dir, "class.png"), "Diagramme de classes")
        if "Base de Données" in title: add_fig(os.path.join(puml_dir, "database_erd.png"), "Schéma Relationnel")
        if "collecte automatique" in title: add_fig(os.path.join(puml_dir, "sequence_scraping.png"), "Pipeline de Scraping")
        if "compatibilité" in title: add_fig(os.path.join(puml_dir, "sequence_compatibility.png"), "Règles de compatibilité")
        if "Interface utilisateur" in title: add_fig(os.path.join(assets_dir, "figure_7_home_1778642725154.png"), "Configurateur PC")

    # Formatting
    for section in doc.sections:
        section.top_margin = Inches(0.98); section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98); section.right_margin = Inches(0.98)
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1.5
        if len(p.text) > 80: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            if not run.bold: run.font.size = Pt(12)

    doc.save(final_path)
    print("Final Indexed Report Built.")

if __name__ == "__main__":
    build_final_indexed_report()
