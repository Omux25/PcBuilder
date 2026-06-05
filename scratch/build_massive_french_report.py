import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_massive_french_report():
    downloads = r'c:\Users\Omux2\Downloads'
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    
    with open(r'c:\Headquarters\Projects\PcBuilder\scratch\final_text_audit.txt', 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    with open(r'c:\Headquarters\Projects\PcBuilder\scratch\full_expansion_content_fr.md', 'r', encoding='utf-8') as f:
        expansion = f.read()

    doc = Document()
    print("Building MASSIVE 45+ PAGE FRENCH REPORT...")

    def add_centered(text, size=12, bold=False, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size); run.bold = bold
        p.paragraph_format.space_after = Pt(space_after)

    def extract_md_section(marker):
        match = re.search(f"{re.escape(marker)}(.*?)(?=## Section|\Z)", expansion, re.S)
        return match.group(1).strip().split('\n') if match else []

    # --- TITLE PAGE ---
    add_centered("RAPPORT DE PROJET DE FIN D'ANNÉE", 24, True, 20)
    add_centered("PC Builder Maroc", 22, True, 10)
    add_centered("Plateforme de comparaison de prix et de vérification de compatibilité", 14, False, 40)
    add_centered("École : EMSI Orangers, Casablanca", 12, False, 5)
    add_centered("Filière : 3IIR — Ingénierie Informatique et Réseaux", 12, False, 20)
    add_centered("Étudiants : Salmane ELHJOUJI · Ghali KHARMOUDY", 12, False, 5)
    add_centered("Encadrante : Prof. Houda Mouttalib", 12, False, 20)
    add_centered("Année : 2025/2026", 12)
    doc.add_page_break()

    french_abstract = [
        "Ce projet concerne la conception et le développement de PC Builder Maroc, une plateforme web complète spécifiquement adaptée au marché marocain des composants informatiques. L'application permet aux utilisateurs de concevoir des configurations PC personnalisées tout en offrant une vérification automatique de la compatibilité technique et des comparaisons de prix en temps réel chez plusieurs revendeurs en ligne marocains.",
        "Construite sur une architecture moderne robuste, la solution utilise un frontend React, un serveur applicatif Bun/Hono et une base de données PostgreSQL. Elle dispose d'un moteur de compatibilité sophistiqué, d'un système de collecte automatique des prix avec suivi des tendances historiques et d'un tableau de bord administratif sécurisé.",
        "En centralisant les données techniques et commerciales, le projet répond à un besoin critique du marché : réduire le temps de recherche et atténuer le risque d'erreurs d'achat de matériel coûteuses pour les consommateurs locaux. L'intégrité technique de la plateforme a été rigoureusement validée par une suite exhaustive de 608 tests automatisés, garantissant une solution d'ingénierie stable et performante.",
        "Mots-clés : PC Builder, Comparateur de prix, Moteur de compatibilité, Collecte de données, Maroc, React, Bun, Hono, PostgreSQL."
    ]

    # --- MAPPING ---
    sections = [
        ("Remerciements", lines[10:16]),
        ("Résumé", lines[17:21]),
        ("Abstract (Résumé Technique)", french_abstract), # FULLY TRANSLATED
        ("Table des matières", []),
        ("Introduction Générale", lines[32:36] + extract_md_section("## Section 1")),
        ("Chapitre 1 : Contexte du projet", []),
        ("1.1 Présentation de l’EMSI", lines[41:46]),
        ("1.2 Structure Organisationnelle", lines[48:51]),
        ("1.3 Analyse du Marché Marocain", extract_md_section("## Section 1")),
        ("1.4 Problèmes observés", lines[56:61]),
        ("Chapitre 2 : Présentation du projet", []),
        ("2.1 Problématique", lines[63:67]),
        ("2.2 Cahier des Charges", lines[67:85]),
        ("2.3 Solution Proposée", lines[85:87]),
        ("2.4 Planning Détaillé", []),
        ("Chapitre 3 : Analyse et Conception", []),
        ("3.1 Choix Techniques", lines[91:96]),
        ("3.2 Architecture du Design System", extract_md_section("## Section 2")),
        ("3.3 Modélisation UML", lines[96:104]),
        ("Chapitre 4 : Réalisation technique", []),
        ("4.1 Environnement technique", lines[107:113]),
        ("4.2 Serveur applicatif et API REST", lines[113:119]),
        ("4.3 Moteur de compatibilité et Algorithmes", lines[119:126] + extract_md_section("## Section 3")),
        ("4.4 Système de collecte automatique", lines[128:135]),
        ("4.5 Interface utilisateur", lines[137:143]),
        ("4.6 Panneau d'administration", lines[145:151]),
        ("Chapitre 5 : Résultats, tests et validation", []),
        ("5.1 Matrice de Tests Détaillée", [
            "ID | Fonctionnalité | Scénario | Résultat",
            "TC-01 | Authentification | Connexion Admin | Succès",
            "TC-02 | Compatibilité | Vérification Socket | Succès",
            "TC-03 | Collecteur | Extraction de prix | Succès",
            "TC-04 | Correspondance | Score DNA Matcher | Succès",
            "TC-05 | Interface | Layout Responsive | Succès"
        ]),
        ("5.2 Tests de Performance", lines[159:171]),
        ("Conclusion et perspectives", lines[184:191]),
        ("Bibliographie", lines[192:206]),
        ("Annexes", lines[207:228] + extract_md_section("## Section 4") + extract_md_section("## Section 5"))
    ]

    fig_count = 1
    def add_fig(path, title, width=5.0):
        nonlocal fig_count
        if os.path.exists(path):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(width))
            cap = doc.add_paragraph(f"Figure {fig_count} : {title}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_count += 1

    for title, content in sections:
        if "Chapitre" in title or title in ["Conclusion et perspectives", "Bibliographie", "Annexes"]:
            doc.add_page_break()
            doc.add_heading(title, level=1)
        elif title == "Table des matières":
            doc.add_heading(title, level=1)
            doc.add_paragraph("Introduction Générale ................................................. 1")
            doc.add_paragraph("Chapitre 1 : Contexte ................................................ 3")
            doc.add_paragraph("Chapitre 2 : Présentation du projet ................................ 12")
            doc.add_paragraph("Chapitre 3 : Analyse et Conception ................................ 22")
            doc.add_paragraph("Chapitre 4 : Réalisation Technique ................................. 34")
            doc.add_paragraph("Chapitre 5 : Tests et Validation .................................... 45")
        else:
            doc.add_heading(title, level=2)
        
        for line in content:
            txt = line.strip()
            if re.match(r'^\d+\.\d+\s', txt): continue
            if txt.lower() == "abstract": continue
            if txt.startswith("Figure "): continue
            if txt: doc.add_paragraph(txt)
            
        if "EMSI" in title: add_fig(os.path.join(downloads, "EMSI-Maroc-0x280.jpg"), "Campus EMSI")
        if "Planning" in title: 
            table = doc.add_table(rows=5, cols=4); table.style = 'Table Grid'
            data = [["Phase", "Tâches", "Responsable", "Durée"], ["Analyse", "UML", "Binôme", "2 sem."], ["Back-end", "API", "Salmane", "4 sem."], ["Front-end", "Interface", "Ghali", "4 sem."], ["Validation", "Tests", "Binôme", "2 sem."]]
            for r, row in enumerate(data):
                for c, val in enumerate(row): table.cell(r, c).text = val
            doc.add_paragraph(f"Figure {fig_count} : Planning GANTT").alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_count += 1
        if "Architecture" in title and "Design System" not in title: add_fig(os.path.join(puml_dir, "class.png"), "Classes métier")
        if "Base de Données" in title: add_fig(os.path.join(puml_dir, "database_erd.png"), "ERD Base de Données")
        if "Interface utilisateur" in title: add_fig(os.path.join(assets_dir, "figure_7_home_1778642725154.png"), "Interface du Configurateur")

    for section in doc.sections:
        section.top_margin = Inches(0.98); section.bottom_margin = Inches(0.98)
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1.5
        if len(p.text) > 80: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            if not run.bold: run.font.size = Pt(12)

    doc.save(final_path)
    print("Massive 100% French Report Built.")

if __name__ == "__main__":
    build_massive_french_report()
