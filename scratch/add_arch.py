import docx
import sys
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client
import os

def add_architecture(doc_path, output_path):
    doc = docx.Document(doc_path)

    # 3. Add Global Architecture Diagram Placeholder
    inserted_arch = False
    
    # We will find the paragraph that contains "Chapitre 3" and insert right after it.
    for i, p in enumerate(doc.paragraphs):
        if "Chapitre 3" in p.text and (p.style.name.startswith('Heading 1') or p.style.name.startswith('Titre 1')):
            # We want to insert after this paragraph.
            # python-docx only has insert_paragraph_before(), so we can do it on the *next* paragraph.
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i+1]
                
                arch_title = next_p.insert_paragraph_before("3.X Schéma d'Architecture Global")
                arch_title.style = 'Heading 2'
                
                arch_body = next_p.insert_paragraph_before("Cette section présente l'architecture technique globale du projet, illustrant les flux de données entre le frontend (React), le backend (Hono/Bun), la base de données (PostgreSQL) et le moteur de scraping.")
                arch_body.style = 'Normal'
                
                arch_box = next_p.insert_paragraph_before()
                r = arch_box.add_run("[PLACEHOLDER IMAGE: Insérer ici un Schéma d'Architecture Global détaillant les interactions techniques du système]")
                r.bold = True
                r.font.color.rgb = RGBColor(255, 102, 0)
                arch_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                inserted_arch = True
                break

    doc.save(output_path)
    print(f"Modifications saved to {output_path}")

def update_tables(doc_path):
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    try:
        abs_path = os.path.abspath(doc_path)
        doc = word.Documents.Open(abs_path)
        
        # Also need to enforce the style on the newly added elements to make sure it's Times New Roman
        for p in doc.Paragraphs:
            if "Schéma d'Architecture Global" in p.Range.Text:
                if "Heading 2" in p.Style.NameLocal or "Titre 2" in p.Style.NameLocal:
                    p.Range.Font.Name = "Times New Roman"
                    p.Range.Font.Size = 14
                    p.Range.Font.Bold = True
                    p.Range.Font.Color = 0
            if "Cette section présente l'architecture" in p.Range.Text:
                p.Range.Font.Name = "Times New Roman"
                p.Range.Font.Size = 12
                p.Format.LineSpacingRule = 1
                p.Format.Alignment = 3 # Justify

        if doc.TablesOfContents.Count > 0:
            for toc in doc.TablesOfContents:
                toc.Update()
        doc.Fields.Update()
        doc.Save()
        print("TOC Updated")
    except Exception as e:
        print(f"Error updating TOC: {e}")
    finally:
        try:
            doc.Close()
        except:
            pass
        word.Quit()

if __name__ == "__main__":
    add_architecture(sys.argv[1], sys.argv[2])
    update_tables(sys.argv[2])
