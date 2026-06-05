import os
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_ultimate_report():
    # 1. SETUP
    original_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport.docx'
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    downloads = r'c:\Users\Omux2\Downloads'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    
    shutil.copy2(original_path, final_path)
    doc = Document(final_path)
    
    # 2. CONTENT EXPANSION (Professional French)
    expansions = {
        "2.1 Problématique": """
Le marché du matériel informatique au Maroc se caractérise par une asymétrie d'information majeure. Un utilisateur souhaitant assembler un ordinateur doit faire face à deux obstacles critiques :
1. La fragmentation de l'offre : Les revendeurs (UltraPC, NextLevel, SetupGame, etc.) opèrent de manière isolée, rendant la comparaison manuelle des prix extrêmement chronophage.
2. La complexité technique : L'interdépendance des composants (compatibilité du socket CPU, puissance de l'alimentation, dimensions du boîtier) crée un risque élevé d'erreurs d'achat coûteuses pour les néophytes.""",
        
        "2.2.1 Besoins Fonctionnels": """
- Pour l'Utilisateur :
    - Recherche et filtrage multicritères par catégorie de composants.
    - Configurateur assisté avec vérification en temps réel de la compatibilité.
    - Comparaison dynamique des prix chez les principaux revendeurs marocains.
    - Visualisation de l'historique des prix pour optimiser le moment de l'achat.
- Pour l'Administrateur :
    - Gestion du catalogue centralisé (CRUD des composants).
    - Supervision des sessions de collecte automatique (Scraping).
    - Mapping manuel des produits non reconnus automatiquement par le système.""",
        
        "3.1 Justification des Choix Techniques": """
Le choix de la stack technologique a été dicté par des impératifs de performance et de modernité.
- React & TypeScript : L'utilisation de React permet une gestion efficace de l'état du configurateur. TypeScript apporte un typage statique crucial.
- Bun & Hono : Bun a été choisi comme "runtime" pour sa rapidité supérieure à Node.js. Hono offre une API légère et extrêmement performante.
- PostgreSQL : Pour la persistence des données, PostgreSQL a été retenu pour sa gestion robuste des relations.""",
        
        "4.1 Environnement technique": """
Le projet est organisé dans un dépôt unique (Monorepo). Le dossier du serveur contient l'API, les services, les routes et les collecteurs. Le dossier de l'interface utilisateur contient l'application publique.
- Scraping Engine : Utilisation de Cheerio pour le parsing DOM.
- DNA Matcher : Algorithme de correspondance de chaînes pour lier les offres au catalogue.
- Compatibility Rules : Moteur de 8 règles métier strictes."""
    }

    # 3. ASSET MAPPING
    assets_map = {
        "Insérer ici une image de la façade du campus EMSI Orangers": os.path.join(downloads, "EMSI-Maroc-0x280.jpg"),
        "Insérer ici l'Organigramme officiel": os.path.join(puml_dir, "emsi_organigramme.png"),
        "3.2 Modélisation UML : Cas d'utilisation": os.path.join(puml_dir, "use_case.png"),
        "3.3 Architecture et Diagramme de Classes": os.path.join(puml_dir, "class.png"),
        "3.4 Conception de la Base de Données": os.path.join(assets_dir, "figure_6_database_erd_1778640799727.png"),
        "Figure 7 : Lecture visuelle du moteur de compatibilité": os.path.join(puml_dir, "sequence_compatibility.png"),
        "Figure 8 : Chaîne de collecte automatique des prix": os.path.join(puml_dir, "sequence_scraping.png"),
        "Figure 9 : Aperçu fonctionnel des interfaces": os.path.join(assets_dir, "figure_7_home_1778642725154.png"),
        "Figure 10 : Synthèse visuelle des résultats": os.path.join(assets_dir, "figure_10_admin_dashboard_1778643047831.png")
    }

    # 4. PROCESSING
    in_mermaid = False
    to_delete = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Mermaid cleaning
        if text.startswith("```mermaid") or text == "```":
            in_mermaid = not in_mermaid
            to_delete.append(i)
            continue
        if in_mermaid:
            to_delete.append(i)
            continue
            
        # Expansion
        for header, content in expansions.items():
            if header in text and len(text) < 150:
                para.text = f"{header}\n{content}"
                break
                
        # Asset insertion
        for placeholder, img_path in assets_map.items():
            if placeholder in text:
                para.text = text.replace(placeholder, "")
                if os.path.exists(img_path):
                    run = para.add_run()
                    run.add_break()
                    run.add_picture(img_path, width=Inches(5.2))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

    # 5. REMOVE MARKED PARAGRAPHS
    for index in sorted(to_delete, reverse=True):
        p = doc.paragraphs[index]
        parent = p._element.getparent()
        parent.remove(p._element)

    # 6. GLOBAL FORMATTING (EMSI GUIDE)
    for section in doc.sections:
        section.top_margin = Inches(0.98) # 2.5cm
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.98)

    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = 1.5
        # Justify main text, Left align lists
        if len(para.text) > 100 and not para.text.strip().startswith(("-", "•")):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Times New Roman
        for run in para.runs:
            run.font.name = 'Times New Roman'
            if not run.bold: run.font.size = Pt(12)

    # 7. ADD ANNEXES AND BIBLIOrevamp
    # (Simplified for one-pass script)
    
    doc.save(final_path)
    print("Ultimate report generation complete.")

if __name__ == "__main__":
    generate_ultimate_report()
