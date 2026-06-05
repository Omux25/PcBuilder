import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import sys
import re

def create_raw_doc(input_text_file, code_files, output_file):
    doc = docx.Document()
    
    with open(input_text_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    paragraphs_text = []
    for line in lines:
        if line.startswith("["):
            match = re.match(r'\[\d+\] .*?: (.*)', line)
            if match:
                text = match.group(1).strip()
                if text:
                    paragraphs_text.append(text)
                    
    iso_690_replacements = {
        "BOEHM, Barry. A spiral model of software development and enhancement. Computer, vol. 21, n° 5, p. 61-72, 1988.": "BOEHM, Barry. A spiral model of software development and enhancement. Computer, 1988, vol. 21, no 5, p. 61-72.",
        "BRIN, Sergey et PAGE, Lawrence. The anatomy of a large-scale hypertextual Web search engine. Computer Networks and ISDN Systems, vol. 30, n° 1-7, p. 107-117, 1998.": "BRIN, Sergey; PAGE, Lawrence. The anatomy of a large-scale hypertextual Web search engine. Computer Networks and ISDN Systems, 1998, vol. 30, no 1-7, p. 107-117.",
        "Claessen, K., & Hughes, J. (2000). QuickCheck: A lightweight tool for random testing of Haskell programs. ACM SIGPLAN Notices, 35(9), 268-279.": "CLAESSEN, Koen; HUGHES, John. QuickCheck: a lightweight tool for random testing of Haskell programs. ACM SIGPLAN Notices, 2000, vol. 35, no 9, p. 268-279.",
        "Codd, E. F. (1970). A relational model of data for large shared data banks. Communications of the ACM, 13(6), 377-387.": "CODD, Edgar F. A relational model of data for large shared data banks. Communications of the ACM, 1970, vol. 13, no 6, p. 377-387.",
        "FIELDING, Roy T. et TAYLOR, Richard N. Principled design of the modern Web architecture. ACM Transactions on Internet Technology, vol. 2, n° 2, p. 115-150, 2002.": "FIELDING, Roy T.; TAYLOR, Richard N. Principled design of the modern Web architecture. ACM Transactions on Internet Technology (TOIT), 2002, vol. 2, no 2, p. 115-150.",
        "Saltzer, J. H., & Schroeder, M. D. (1975). The protection of information in computer systems. Proceedings of the IEEE, 63(9), 1278-1308.": "SALTZER, Jerome H.; SCHROEDER, Michael D. The protection of information in computer systems. Proceedings of the IEEE, 1975, vol. 63, no 9, p. 1278-1308.",
        "Jones, M., Bradley, J., & Sakimura, N. (2015). JSON Web Token (JWT) (RFC 7519). Internet Engineering Task Force.": "JONES, Michael; BRADLEY, John; SAKIMURA, Nat. JSON Web Token (JWT). RFC 7519. Internet Engineering Task Force, 2015.",
        "OWASP Foundation. (2021). OWASP Top 10:2021. https://owasp.org/Top10/": "OWASP FOUNDATION. OWASP Top 10:2021 [en ligne]. 2021. Disponible sur : https://owasp.org/Top10/ (consulté le 13 mai 2026).",
        "POSTGRESQL GLOBAL DEVELOPMENT GROUP. PostgreSQL 16 Documentation [en ligne]. 2025. Disponible sur : https://www.postgresql.org/docs/16/": "POSTGRESQL GLOBAL DEVELOPMENT GROUP. PostgreSQL 16 Documentation [en ligne]. 2025. Disponible sur : https://www.postgresql.org/docs/16/ (consulté le 13 mai 2026).",
        "React Team. (2026). React documentation. https://react.dev/": "REACT TEAM. React documentation [en ligne]. 2026. Disponible sur : https://react.dev/ (consulté le 13 mai 2026).",
        "Bun Team. (2026). Bun documentation. https://bun.sh/docs": "BUN TEAM. Bun documentation [en ligne]. 2026. Disponible sur : https://bun.sh/docs (consulté le 13 mai 2026).",
        "Hono Team. (2026). Hono documentation. https://hono.dev/": "HONO TEAM. Hono documentation [en ligne]. 2026. Disponible sur : https://hono.dev/ (consulté le 13 mai 2026).",
        "PCPartPicker. (2026). PCPartPicker. https://pcpartpicker.com/": "PCPARTPICKER. PCPartPicker [en ligne]. 2026. Disponible sur : https://pcpartpicker.com/ (consulté le 13 mai 2026)."
    }

    placeholders = {
        "(Insérer Figure 1 : Campus EMSI Orangers ici)": "[PLACEHOLDER IMAGE: Photo officielle du Campus EMSI Orangers]",
        "(Insérer Figure 2 : Organigramme ici)": "[PLACEHOLDER IMAGE: Organigramme de l'équipe du projet]",
        "(Insérer Figure 4 : Diagramme de cas d'utilisation ici)": "[PLACEHOLDER IMAGE: Diagramme de Cas d'Utilisation (UML)]",
        "(Insérer Figure 5 : Diagramme de classes métier ici)": "[PLACEHOLDER IMAGE: Diagramme de Classes (UML)]",
        "(Insérer Figure 6 : Schéma Relationnel / ERD ici)": "[PLACEHOLDER IMAGE: Schéma Relationnel de la Base de Données (MCD/MLD)]",
        "(Insérer Figure 7 : Règles de compatibilité ici)": "[PLACEHOLDER IMAGE: Logigramme illustrant le Moteur de Compatibilité]",
        "(Insérer Figure 8 : Pipeline de Scraping ici)": "[PLACEHOLDER IMAGE: Schéma d'Architecture du Pipeline de Scraping]",
        "(Insérer Figure 9 : UI Configurateur PC ici)": "[PLACEHOLDER IMAGE: Capture d'écran 1 : Le Configurateur PC]",
        "Annexe B : Extrait du modèle relationnel (Insérer l'image du modèle de base de données ici)": "[PLACEHOLDER IMAGE: Capture d'écran détaillée du Schéma de Base de données]",
        "[Insérer Logo EMSI ici]": "[PLACEHOLDER LOGO: Logo officiel de l'EMSI]"
    }
    
    captions = {
        "(Insérer Figure 1 : Campus EMSI Orangers ici)": "Figure 1 : Campus EMSI Orangers",
        "(Insérer Figure 2 : Organigramme ici)": "Figure 2 : Organigramme",
        "(Insérer Figure 4 : Diagramme de cas d'utilisation ici)": "Figure 4 : Diagramme de cas d'utilisation",
        "(Insérer Figure 5 : Diagramme de classes métier ici)": "Figure 5 : Diagramme de classes métier",
        "(Insérer Figure 6 : Schéma Relationnel / ERD ici)": "Figure 6 : Schéma Relationnel / ERD",
        "(Insérer Figure 7 : Règles de compatibilité ici)": "Figure 7 : Règles de compatibilité",
        "(Insérer Figure 8 : Pipeline de Scraping ici)": "Figure 8 : Pipeline de Scraping",
        "(Insérer Figure 9 : UI Configurateur PC ici)": "Figure 9 : Aperçu fonctionnel des interfaces et Configurateur PC"
    }

    additional_placeholders = [
        ("[PLACEHOLDER IMAGE: Capture d'écran 2 : Comparateur de prix]", "Figure 10 : Comparateur de Prix multicritères"),
        ("[PLACEHOLDER IMAGE: Capture d'écran 3 : Page de Tendances]", "Figure 11 : Page des Tendances et de l'Historique")
    ]

    page_break_before = [
        "Remerciements", "Résumé", "Abstract", "Table des matières", 
        "Liste des figures", "Introduction Générale", 
        "Chapitre 1 : Contexte du projet", "Chapitre 2 : Présentation du projet",
        "Chapitre 3 : Analyse et Conception", "Chapitre 4 : Réalisation technique",
        "Chapitre 5 : Résultats, tests et validation", "Conclusion et perspectives",
        "Bibliographie", "Annexes"
    ]
    
    h2_prefixes = [
        "1.1", "1.2", "1.3", "1.4",
        "2.1", "2.2", "2.3", "2.4",
        "3.1", "3.2", "3.3", "3.4",
        "4.1", "4.2", "4.3", "4.4", "4.5", "4.6",
        "5.1", "5.2", "5.3", "5.4", "5.5", "5.6"
    ]
    
    is_cover_page = True
    skip_next = False

    for i, text in enumerate(paragraphs_text):
        if skip_next:
            skip_next = False
            continue

        if text in iso_690_replacements:
            text = iso_690_replacements[text]

        is_h1 = text in page_break_before
        is_h2 = any(text.startswith(prefix + " ") for prefix in h2_prefixes)
        
        if is_cover_page and text == "Remerciements":
            is_cover_page = False
            
        if is_h1 and not is_cover_page:
            doc.add_page_break()
            doc.add_paragraph(text, style='Heading 1')
            
            if text == "Table des matières" or text == "Liste des figures":
                doc.add_paragraph() # Leave a blank para for win32com to inject into
                skip_next = True
            continue
            
        if is_h2:
            doc.add_paragraph(text, style='Heading 2')
            if text.startswith("5.2 Tests"):
                doc.add_paragraph("Le projet compte 608 tests automatisés : tests unitaires, tests d'intégration et tests basés sur les propriétés. Le tableau suivant synthétise les résultats obtenus :")
                
                table = doc.add_table(rows=5, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Type de Test'
                hdr[1].text = 'Nombre Total'
                hdr[2].text = 'Taux de Réussite'
                hdr[3].text = 'Couverture Principale'
                
                data = [
                    ("Tests Unitaires", "412", "98%", "Règles de compatibilité, utilitaires, formatage"),
                    ("Tests d'Intégration", "154", "100%", "Routes API, accès DB, services métier"),
                    ("Tests de Propriétés (Fast-check)", "42", "100%", "Validation Zod, pagination, algorithmes de matching"),
                    ("Total / Moyenne", "608", "~99%", "87% de la logique métier globale")
                ]
                for row_idx, row_data in enumerate(data):
                    cells = table.rows[row_idx+1].cells
                    cells[0].text = row_data[0]
                    cells[1].text = row_data[1]
                    cells[2].text = row_data[2]
                    cells[3].text = row_data[3]
                    
                skip_next = True
            continue
            
        if text.startswith("RAPPORT DE PROJET DE FIN D'ANNÉE") and is_cover_page:
            p_logo = doc.add_paragraph(placeholders["[Insérer Logo EMSI ici]"])
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p3 = doc.add_paragraph("\n\nRAPPORT DE PROJET DE FIN D'ANNÉE")
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        if "PC Builder Maroc Plateforme de comparaison" in text or "PC Builder Maroc" in text and is_cover_page:
            p = doc.add_paragraph("PC Builder Maroc\nPlateforme de comparaison de prix et de vérification de compatibilité\n\n\n\n")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        if is_cover_page:
            p = doc.add_paragraph("École : EMSI Orangers, Casablanca\n"
                                  "Filière : 3IIR — Ingénierie Informatique et Réseaux\n\n"
                                  "Réalisé par : Salmane ELHJOUJI et Ghali KHARMOUDY\n"
                                  "Encadrante : Prof. Houda Mouttalib\n\n"
                                  "Année universitaire : 2025/2026\n"
                                  "Date de soutenance : 14 Mai 2026\n")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        if text == "(Insérer Figure 3 : Planning ici)":
            doc.add_paragraph("Figure 3 : Planning du projet (Gantt/Tableau)")
            table = doc.add_table(rows=5, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Phase'
            hdr_cells[1].text = 'Tâches'
            hdr_cells[2].text = 'Responsable'
            hdr_cells[3].text = 'Durée'
            
            data = [
                ("Analyse", "UML, Spécifications", "Binôme", "2 semaines"),
                ("Back-end", "API, Scraping, DB", "Salmane", "4 semaines"),
                ("Front-end", "UI React, Configurateur", "Ghali", "4 semaines"),
                ("Validation", "Tests, Déploiement", "Binôme", "2 semaines")
            ]
            for r_idx, row in enumerate(data):
                cells = table.rows[r_idx+1].cells
                cells[0].text = row[0]
                cells[1].text = row[1]
                cells[2].text = row[2]
                cells[3].text = row[3]
            continue
            
        if text in placeholders:
            box = doc.add_paragraph(placeholders[text])
            box.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if text in captions:
                cap = doc.add_paragraph(captions[text])
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if text == "(Insérer Figure 9 : UI Configurateur PC ici)":
                for additional, add_cap in additional_placeholders:
                    b2 = doc.add_paragraph(additional)
                    b2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    cap2 = doc.add_paragraph(add_cap)
                    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
                
        if text == "Annexe D : Résumé des tests automatisés" or text in ["Total Tests : 608", "Tests Unitaires : 412 (98% pass)", "Tests d'Intégration : 154 (100% pass)", "Tests de Propriétés (Fast-check) : 42 (100% pass)", "Couverture globale : 87% de la logique métier.", "(À générer automatiquement avec Word)"]:
            continue

        doc.add_paragraph(text)
            
    doc.add_page_break()
    doc.add_paragraph("Annexe F : Extraits de Code Source Majeurs", style='Heading 1')
    
    for c_file in code_files:
        try:
            with open(c_file, 'r', encoding='utf-8') as cf:
                content = cf.read()
            doc.add_paragraph(f"Fichier : {c_file.split('/')[-1]}", style='Heading 2')
            doc.add_paragraph(content)
        except:
            pass
            
    doc.save(output_file)

if __name__ == "__main__":
    code_files = [
        "c:/Headquarters/Projects/PcBuilder/shared/compatibility-engine.ts",
        "c:/Headquarters/Projects/PcBuilder/shared/types.ts"
    ]
    create_raw_doc(sys.argv[1], code_files, sys.argv[2])
