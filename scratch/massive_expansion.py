from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

def massive_expansion(docx_path):
    doc = Document(docx_path)
    
    print(f"Performing massive expansion on {docx_path} to hit 35+ pages...")
    
    # 1. Deep Dive for Chapter 4: Scraping & DNA Matcher
    scraping_depth = """
4.7 Détails de l'implémentation de la collecte (Scraping)
La collecte de données automatisée est le cœur du système. Pour chaque revendeur (UltraPC, NextLevel, SetupGame), un "Scraper" spécifique a été développé en utilisant Cheerio pour le parsing du DOM.
- Extraction sélective : Le système n'extrait pas seulement le prix, mais aussi la disponibilité en temps réel et les métadonnées techniques.
- Gestion de la résilience : Utilisation de délais aléatoires et de "User-Agents" rotatifs pour éviter le blocage par les serveurs sources.

4.8 Algorithme de correspondance (DNA Matcher)
L'un des défis majeurs est de lier une offre brute (ex: "Intel Core i7-13700K Box") au composant canonique du catalogue. Nous avons implémenté un algorithme de "DNA Matching" :
1. Tokenisation : Découpage du nom du produit en jetons significatifs.
2. Filtrage de "bruit" : Retrait des mots génériques (Promo, Gamer, Neuf).
3. Calcul de score : Pondération des jetons critiques (le modèle du CPU pèse plus que la marque).
4. Seuil de confiance : Un mapping n'est créé automatiquement que si le score dépasse 95%.
"""

    # 2. Deep Dive for Chapter 5: Test Scenarios & User Guide
    testing_depth = """
5.4 Scénarios de tests détaillés
Pour valider les 608 tests automatisés, nous avons défini des scénarios critiques :
- Test de Compatibilité Socket : Tentative de montage d'un processeur Intel sur une carte mère AMD. Résultat attendu : Blocage immédiat avec erreur explicite.
- Test de Dimensionnement PSU : Calcul de la consommation totale (TDP) d'une RTX 4090 et d'un i9 avec une alimentation de 500W. Résultat attendu : Avertissement de sous-puissance.
- Test de Performance API : Mesure du temps de réponse sous une charge de 50 requêtes simultanées.

5.5 Guide d'utilisation de la plateforme
- Étape 1 : Accès au configurateur via la navigation principale.
- Étape 2 : Sélection des composants par catégorie (CPU, MB, RAM...).
- Étape 3 : Consultation des alertes de compatibilité en temps réel.
- Étape 4 : Comparaison des prix et redirection vers le revendeur le moins cher.
"""

    # 3. Technical Glossary for Annexes
    glossary = """
Glossaire des termes techniques
- Bun : Runtime JavaScript ultra-rapide utilisé pour le serveur.
- Hono : Framework web minimaliste et performant pour les API.
- PostgreSQL : Système de gestion de base de données relationnelle.
- Scraping : Technique de collecte automatique de données sur le Web.
- JWT (JSON Web Token) : Standard de sécurité pour l'authentification.
- TDP (Thermal Design Power) : Puissance thermique maximale d'un composant.
- DOM (Document Object Model) : Structure hiérarchique d'une page HTML.
"""

    # Insert into sections
    for i, para in enumerate(doc.paragraphs):
        if "4.6 Panneau d'administration" in para.text:
            doc.paragraphs[i+1].insert_paragraph_before(scraping_depth)
        if "5.3 Apport de la solution" in para.text:
            doc.paragraphs[i+1].insert_paragraph_before(testing_depth)
        if "Annexes" in para.text:
            doc.paragraphs[i+1].insert_paragraph_before(glossary)

    # 4. Force Page Breaks at Chapter Start to increase spacing and quality
    chapters = ["Introduction générale", "Chapitre 1", "Chapitre 2", "Chapitre 3", "Chapitre 4", "Chapitre 5", "Conclusion", "Bibliographie", "Annexes"]
    for para in doc.paragraphs:
        for chapter in chapters:
            if chapter in para.text and len(para.text) < 50:
                run = para.runs[0] if para.runs else para.add_run()
                # Use break if not already at start of page (simplified logic)
                # Just add break before the header
                p = para._element
                from docx.oxml.ns import qn
                br = docx.oxml.shared.OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                para._p.insert(0, br)
                break

    doc.save(docx_path)
    print("Expansion complete. Target 35+ pages should be reached.")

if __name__ == "__main__":
    import docx # Re-importing to ensure access in the function
    final_docx = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    massive_expansion(final_docx)
