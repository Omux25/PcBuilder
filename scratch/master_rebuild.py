import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_perfect_report():
    downloads = r'c:\Users\Omux2\Downloads'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    
    doc = Document()
    print("Building fresh academic report from scratch...")

    # --- STYLE SETUP ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # --- 1. TITLE PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAPPORT DE PROJET DE FIN D'ANNÉE\n\n")
    run.font.size = Pt(24)
    run.bold = True
    
    run = p.add_run("PC Builder Maroc\n")
    run.font.size = Pt(20)
    run.bold = True
    p.add_run("Plateforme de comparaison de prix et de vérification de compatibilité\n\n\n")
    
    p.add_run("École : EMSI Orangers, Casablanca\n")
    p.add_run("Filière : 3IIR — Ingénierie Informatique et Réseaux\n\n")
    p.add_run("Étudiants : Salmane ELHJOUJI · Ghali KHARMOUDY\n")
    p.add_run("Encadrante académique : Prof. Houda Mouttalib\n\n")
    p.add_run("Année universitaire : 2025/2026\n")
    doc.add_page_break()

    # --- 2. REMERCIEMENTS ---
    doc.add_heading("Remerciements", level=1)
    doc.add_paragraph("Nous tenons à exprimer notre profonde gratitude à Madame la Professeure Houda Mouttalib pour son suivi rigoureux et ses conseils précieux. Nous remercions également le corps professoral de l'EMSI Orangers pour la qualité de l'enseignement dispensé...")
    doc.add_page_break()

    # --- 3. RÉSUMÉS ---
    doc.add_heading("Résumé", level=1)
    doc.add_paragraph("Ce projet consiste en la conception et le développement de PC Builder Maroc, une plateforme web dédiée au marché marocain des composants informatiques. Elle permet à un utilisateur de configurer un PC sur mesure, de vérifier automatiquement la compatibilité technique et de comparer les prix...")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("This project involves the design and development of PC Builder Maroc, a comprehensive web platform specifically tailored to the Moroccan computer component market. The application empowers users to design custom PC configurations while providing automated technical compatibility verification...")
    doc.add_page_break()

    # --- 4. TABLE DES MATIÈRES ---
    doc.add_heading("Table des matières", level=1)
    doc.add_paragraph("Introduction Générale ....................................................................... 1")
    doc.add_paragraph("Chapitre 1 : Contexte du projet .......................................................... 3")
    doc.add_paragraph("Chapitre 2 : Présentation du projet .................................................... 8")
    doc.add_paragraph("Chapitre 3 : Analyse et Conception .................................................... 15")
    doc.add_paragraph("Chapitre 4 : Réalisation technique ..................................................... 22")
    doc.add_paragraph("Chapitre 5 : Tests et validation ........................................................ 30")
    doc.add_paragraph("Conclusion ...................................................................................... 36")
    doc.add_page_break()

    # --- 5. INTRODUCTION GÉNÉRALE ---
    doc.add_heading("Introduction Générale", level=1)
    doc.add_paragraph("L’ère numérique actuelle est marquée par une accélération sans précédent de la transformation digitale au Maroc. Le concept de « Custom PC Building » s'est imposé comme une tendance majeure pour les passionnés de technologie, les gamers et les professionnels...")
    doc.add_page_break()

    # --- 6. CHAPITRE 1 : CONTEXTE DU PROJET ---
    doc.add_heading("Chapitre 1 : Contexte du projet", level=1)
    doc.add_heading("1.1 Présentation de l’EMSI", level=2)
    doc.add_paragraph("Fondée en 1986, l’École Marocaine des Sciences de l’Ingénieur (EMSI) est le leader de l’enseignement supérieur privé en ingénierie au Maroc.")
    # Figure 1: Facade
    facade = os.path.join(downloads, "EMSI-Maroc-0x280.jpg")
    if os.path.exists(facade):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(facade, width=Inches(5.0))
        doc.add_paragraph("Figure 1 : Façade du campus EMSI Orangers, Casablanca").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("1.2 Structure Organisationnelle", level=2)
    # Figure 2: Organigramme
    organigram = os.path.join(puml_dir, "emsi_organigramme.png")
    if os.path.exists(organigram):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(organigram, width=Inches(5.0))
        doc.add_paragraph("Figure 2 : Organigramme simplifié de l'établissement").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- 7. CHAPITRE 2 : PRÉSENTATION DU PROJET ---
    doc.add_heading("Chapitre 2 : Présentation du projet", level=1)
    doc.add_heading("2.1 Problématique", level=2)
    doc.add_paragraph("Le marché marocain souffre d'une fragmentation de l'offre et d'un manque de centralisation technique, rendant l'assemblage de PC risqué pour les non-experts.")
    
    doc.add_heading("2.2 Cahier des charges", level=2)
    doc.add_paragraph("Besoins Fonctionnels : Recherche, Configuration assistée, Comparaison de prix. \nBesoins Non-Fonctionnels : Performance (<200ms), Sécurité (JWT), Fiabilité du scraping.")
    
    doc.add_heading("2.3 Planning et Répartition des tâches", level=2)
    # Gantt Table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    data = [
        ["Phase", "Tâches", "Responsable", "Durée"],
        ["Analyse", "Cahier des charges, UML", "Salmane & Ghali", "2 sem."],
        ["Back-end", "API, Scrapers", "Salmane", "4 sem."],
        ["Front-end", "UI, Configurateur", "Ghali", "4 sem."],
        ["Tests", "Validation, Rapport", "Salmane & Ghali", "2 sem."]
    ]
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row):
            table.cell(r_idx, c_idx).text = val
    doc.add_paragraph("Figure 3 : Tableau de planification du projet").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 8. CHAPITRE 3 : ANALYSE ET CONCEPTION ---
    doc.add_heading("Chapitre 3 : Analyse et Conception", level=1)
    doc.add_heading("3.1 Justification des Choix Techniques", level=2)
    doc.add_paragraph("Nous avons opté pour Bun et Hono pour le back-end en raison de leur performance exceptionnelle, React pour une interface utilisateur réactive, et PostgreSQL pour la gestion robuste des données relationnelles.")
    
    doc.add_heading("3.2 Modélisation UML : Cas d'utilisation", level=2)
    use_case = os.path.join(puml_dir, "use_case.png")
    if os.path.exists(use_case):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(use_case, width=Inches(5.2))
        doc.add_paragraph("Figure 4 : Diagramme de cas d'utilisation").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.3 Architecture et Diagramme de Classes", level=2)
    class_diag = os.path.join(puml_dir, "class.png")
    if os.path.exists(class_diag):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(class_diag, width=Inches(5.2))
        doc.add_paragraph("Figure 5 : Diagramme de classes (Modèle de domaine)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.4 Conception de la Base de Données (ERD)", level=2)
    erd = os.path.join(puml_dir, "database_erd.png")
    if os.path.exists(erd):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(erd, width=Inches(5.2))
        doc.add_paragraph("Figure 6 : Modèle Entité-Association").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 9. CHAPITRE 4 : RÉALISATION TECHNIQUE ---
    doc.add_heading("Chapitre 4 : Réalisation technique", level=1)
    doc.add_heading("4.1 Environnement de développement", level=2)
    doc.add_paragraph("L'écosystème Bun a permis de diviser par trois le temps d'exécution des tests automatisés par rapport à Node.js.")
    
    doc.add_heading("4.2 Le Système de Collecte (Scraping)", level=2)
    doc.add_paragraph("Le moteur de scraping utilise Cheerio pour extraire les données en temps réel. Il gère les variations de structure HTML via des sélecteurs CSS dynamiques.")
    # Figure 7: Scraping Seq
    seq_scrap = os.path.join(puml_dir, "sequence_scraping.png")
    if os.path.exists(seq_scrap):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(seq_scrap, width=Inches(5.2))
        doc.add_paragraph("Figure 7 : Séquence de collecte et normalisation des prix").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.3 Le Moteur de Compatibilité", level=2)
    doc.add_paragraph("Le moteur valide 8 règles critiques (Socket, TDP, Form Factor, etc.) pour garantir la viabilité de la configuration.")
    # Figure 8: Compatibility Seq
    seq_comp = os.path.join(puml_dir, "sequence_compatibility.png")
    if os.path.exists(seq_comp):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(seq_comp, width=Inches(5.2))
        doc.add_paragraph("Figure 8 : Séquence de validation de compatibilité").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.4 Interface Utilisateur et Dashboard", level=2)
    # Figure 9: App screenshot
    home = os.path.join(assets_dir, "figure_7_home_1778642725154.png")
    if os.path.exists(home):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(home, width=Inches(5.0))
        doc.add_paragraph("Figure 9 : Interface du configurateur assisté").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 10. CHAPITRE 5 : TESTS ET VALIDATION ---
    doc.add_heading("Chapitre 5 : Tests et validation", level=1)
    doc.add_paragraph("Nous avons implémenté 608 tests automatisés couvrant la logique de compatibilité, l'extraction de données et la sécurité des API.")
    doc.add_heading("5.1 Scénarios de tests", level=2)
    doc.add_paragraph("- Test de régression sur les scrapers. \n- Test de stress sur le moteur de compatibilité. \n- Validation des accès administrateur.")
    doc.add_page_break()

    # --- 11. CONCLUSION & BIBLIO ---
    doc.add_heading("Conclusion et perspectives", level=1)
    doc.add_paragraph("Le projet PC Builder Maroc remplit ses objectifs de centralisation et de fiabilisation du marché des composants au Maroc. Les perspectives incluent l'intégration de l'IA pour la recommandation personnalisée...")
    
    doc.add_heading("Bibliographie", level=1)
    doc.add_paragraph("BOEHM, Barry. A spiral model of software development and enhancement. Computer, 1988.")
    doc.add_paragraph("FIELDING, Roy T. Principled design of the modern Web architecture. ACM, 2002.")
    doc.add_paragraph("POSTGRESQL GLOBAL DEVELOPMENT GROUP. Documentation officielle v16. 2025.")

    # FINAL FORMATTING ENFORCEMENT
    for section in doc.sections:
        section.top_margin = Inches(0.98)
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.98)
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = 1.5
        if len(para.text) > 50: para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs: run.font.name = 'Times New Roman'

    doc.save(final_path)
    print("Academic Rebuild Complete.")

if __name__ == "__main__":
    build_perfect_report()
