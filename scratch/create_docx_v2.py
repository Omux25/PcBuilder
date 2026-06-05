import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import sys
import re

def create_styled_doc(input_text_file, code_files, output_file):
    doc = docx.Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        
    # Styles setup
    styles = doc.styles
    
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(12)
    
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(10)

    # Read the raw text
    with open(input_text_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    paragraphs_text = []
    for line in lines:
        if line.startswith("["):
            match = re.match(r'\[\d+\] .*?: (.*)', line)
            if match:
                text = match.group(1).strip()
                if text:
                    paragraphs_text.append(text)
                    
    # ISO 690 Bibliography Adjustments mapping
    iso_690_replacements = {
        "BOEHM, Barry. A spiral model of software development and enhancement. Computer, vol. 21, n° 5, p. 61-72, 1988.": "BOEHM, Barry. A spiral model of software development and enhancement. Computer, 1988, vol. 21, no 5, p. 61-72.",
        "BRIN, Sergey et PAGE, Lawrence. The anatomy of a large-scale hypertextual Web search engine. Computer Networks and ISDN Systems, vol. 30, n° 1-7, p. 107-117, 1998.": "BRIN, Sergey; PAGE, Lawrence. The anatomy of a large-scale hypertextual Web search engine. Computer Networks and ISDN Systems, 1998, vol. 30, no 1-7, p. 107-117.",
        "Claessen, K., & Hughes, J. (2000). QuickCheck: A lightweight tool for random testing of Haskell programs. ACM SIGPLAN Notices, 35(9), 268-279.": "CLAESSEN, Koen; HUGHES, John. QuickCheck: a lightweight tool for random testing of Haskell programs. ACM SIGPLAN Notices, 2000, vol. 35, no 9, p. 268-279.",
        "Codd, E. F. (1970). A relational model of data for large shared data banks. Communications of the ACM, 13(6), 377-387.": "CODD, Edgar F. A relational model of data for large shared data banks. Communications of the ACM, 1970, vol. 13, no 6, p. 377-387.",
        "FIELDING, Roy T. et TAYLOR, Richard N. Principled design of the modern Web architecture. ACM Transactions on Internet Technology, vol. 2, n° 2, p. 115-150, 2002.": "FIELDING, Roy T.; TAYLOR, Richard N. Principled design of the modern Web architecture. ACM Transactions on Internet Technology (TOIT), 2002, vol. 2, no 2, p. 115-150.",
        "Saltzer, J. H., & Schroeder, M. D. (1975). The protection of information in computer systems. Proceedings of the IEEE, 63(9), 1278-1308.": "SALTZER, Jerome H.; SCHROEDER, Michael D. The protection of information in computer systems. Proceedings of the IEEE, 1975, vol. 63, no 9, p. 1278-1308.",
        "Jones, M., Bradley, J., & Sakimura, N. (2015). JSON Web Token (JWT) (RFC 7519). Internet Engineering Task Force.": "JONES, Michael; BRADLEY, John; SAKIMURA, Nat. JSON Web Token (JWT). RFC 7519. Internet Engineering Task Force, 2015.",
        "OWASP Foundation. (2021). OWASP Top 10:2021. https://owasp.org/Top10/": "OWASP FOUNDATION. OWASP Top 10:2021 [en ligne]. 2021. Disponible sur : https://owasp.org/Top10/ (consulté le 13 mai 2026).",
        "POSTGRESQL GLOBAL DEVELOPMENT GROUP. PostgreSQL 16 Documentation [en ligne]. 2025. Disponible sur : https://www.postgresql.org/docs/16/": "POSTGRESQL GLOBAL DEVELOPMENT GROUP. PostgreSQL 16 Documentation [en ligne]. 2025. Disponible sur : https://www.postgresql.org/docs/16/ (consulté le 13 mai 2026).",
        "React Team. (2026). React documentation. https://react.dev/": "REACT TEAM. React documentation [en ligne]. 2026. Disponible sur : https://react.dev/ (consulté le 13 mai 2026).",
        "Bun Team. (2026). Bun documentation. https://bun.sh/docs": "BUN TEAM. Bun documentation [en ligne]. 2026. Disponible sur : https://bun.sh/docs (consulté le 13 mai 2026).",
        "Hono Team. (2026). Hono documentation. https://hono.dev/": "HONO TEAM. Hono documentation [en ligne]. 2026. Disponible sur : https://hono.dev/ (consulté le 13 mai 2026).",
        "PCPartPicker. (2026). PCPartPicker. https://pcpartpicker.com/": "PCPARTPICKER. PCPartPicker [en ligne]. 2026. Disponible sur : https://pcpartpicker.com/ (consulté le 13 mai 2026)."
    }

    placeholders = {
        "(Insérer Figure 1 : Campus EMSI Orangers ici)": "[PLACEHOLDER IMAGE: Photo officielle du Campus EMSI Orangers (Haute résolution). S'assurer que le logo de l'école est visible ou que le bâtiment est clair.]",
        "(Insérer Figure 2 : Organigramme ici)": "[PLACEHOLDER IMAGE: Organigramme de l'équipe du projet ou de la structure. Indiquer clairement Salmane et Ghali et leurs rôles (ex: Backend/Frontend) ainsi que le rôle du tuteur.]",
        "(Insérer Figure 4 : Diagramme de cas d'utilisation ici)": "[PLACEHOLDER IMAGE: Diagramme de Cas d'Utilisation (UML). Doit montrer clairement l'Acteur Utilisateur (Configure un PC, Compare les prix, Voit les tendances) et l'Acteur Admin (Gère le catalogue, Supervise le scraping).]",
        "(Insérer Figure 5 : Diagramme de classes métier ici)": "[PLACEHOLDER IMAGE: Diagramme de Classes (UML). Doit illustrer les classes principales : Composant (avec ses attributs techniques JSONB), Revendeur, Offre, et HistoriquePrix.]",
        "(Insérer Figure 6 : Schéma Relationnel / ERD ici)": "[PLACEHOLDER IMAGE: Schéma Relationnel de la Base de Données (MCD/MLD). Montrer les tables PostgreSQL, les clés étrangères, et mettre en évidence le champ 'specs' de type JSONB.]",
        "(Insérer Figure 7 : Règles de compatibilité ici)": "[PLACEHOLDER IMAGE: Diagramme de Séquence ou Logigramme illustrant le Moteur de Compatibilité. Montrer le flux : Validation Socket -> Validation RAM -> Validation Alimentation (TDP) -> Retour Erreurs/Warnings.]",
        "(Insérer Figure 8 : Pipeline de Scraping ici)": "[PLACEHOLDER IMAGE: Schéma d'Architecture du Pipeline de Scraping. Montrer l'extraction depuis UltraPC/NextLevel/SetupGame, l'algorithme 'DNA Matcher' pour la normalisation, et l'insertion dans PostgreSQL.]",
        "(Insérer Figure 9 : UI Configurateur PC ici)": "[PLACEHOLDER IMAGE: Capture d'écran 1 : Le Configurateur PC. Montrer une configuration en cours avec des alertes de compatibilité (ex: Socket mismatch) générées par le moteur en temps réel.]",
        "Annexe B : Extrait du modèle relationnel (Insérer l'image du modèle de base de données ici)": "[PLACEHOLDER IMAGE: Capture d'écran détaillée du Schéma de Base de données (plus grande que la Figure 6) ou un extrait complexe du code DDL PostgreSQL.]",
        "[Insérer Logo EMSI ici]": "[PLACEHOLDER LOGO: Insérer ici le Logo officiel de l'EMSI en haute qualité avec un fond transparent.]"
    }

    additional_placeholders = [
        "[PLACEHOLDER IMAGE: Capture d'écran 2 : Comparateur de prix. Mettre en évidence un même composant avec ses prix différents chez UltraPC, NextLevel et SetupGame pour montrer la valeur ajoutée du système.]",
        "[PLACEHOLDER IMAGE: Capture d'écran 3 : Page de Tendances. Montrer l'évolution et l'historique d'un prix dans le temps sous forme de graphique analytique.]"
    ]

    page_break_before = [
        "Remerciements", "Résumé", "Abstract", "Table des matières", 
        "Liste des figures", "Introduction Générale", 
        "Chapitre 1 : Contexte du projet", "Chapitre 2 : Présentation du projet",
        "Chapitre 3 : Analyse et Conception", "Chapitre 4 : Réalisation technique",
        "Chapitre 5 : Résultats, tests et validation", "Conclusion et perspectives",
        "Bibliographie", "Annexes"
    ]
    
    h2_prefixes = [
        "1.1", "1.2", "1.3", "1.4",
        "2.1", "2.2", "2.3", "2.4",
        "3.1", "3.2", "3.3", "3.4",
        "4.1", "4.2", "4.3", "4.4", "4.5", "4.6",
        "5.1", "5.2", "5.3", "5.4", "5.5", "5.6"
    ]
    
    is_cover_page = True
    skip_next = False

    for i, text in enumerate(paragraphs_text):
        if skip_next:
            skip_next = False
            continue

        if text in iso_690_replacements:
            text = iso_690_replacements[text]

        is_h1 = text in page_break_before
        is_h2 = any(text.startswith(prefix + " ") for prefix in h2_prefixes)
        
        if is_cover_page and text == "Remerciements":
            is_cover_page = False
            
        if is_h1 and not is_cover_page:
            doc.add_page_break()
            p = doc.add_paragraph(text, style='Heading 1')
            continue
            
        if is_h2:
            p = doc.add_paragraph(text, style='Heading 2')
            
            # Special logic for 5.2 Tests
            if text.startswith("5.2 Tests"):
                # Write the text then immediately the table
                doc.add_paragraph("Le projet compte 608 tests automatisés : tests unitaires, tests d'intégration et tests basés sur les propriétés. Le tableau suivant synthétise les résultats obtenus :")
                
                table = doc.add_table(rows=5, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Type de Test'
                hdr[1].text = 'Nombre Total'
                hdr[2].text = 'Taux de Réussite'
                hdr[3].text = 'Couverture Principale'
                
                data = [
                    ("Tests Unitaires", "412", "98%", "Règles de compatibilité, utilitaires, formatage"),
                    ("Tests d'Intégration", "154", "100%", "Routes API, accès DB, services métier"),
                    ("Tests de Propriétés (Fast-check)", "42", "100%", "Validation Zod, pagination, algorithmes de matching"),
                    ("Total / Moyenne", "608", "~99%", "87% de la logique métier globale")
                ]
                for row_idx, row_data in enumerate(data):
                    cells = table.rows[row_idx+1].cells
                    cells[0].text = row_data[0]
                    cells[1].text = row_data[1]
                    cells[2].text = row_data[2]
                    cells[3].text = row_data[3]
                    
                # We skip the next paragraph because we replaced it
                skip_next = True
            continue
            
        if text.startswith("RAPPORT DE PROJET DE FIN D'ANNÉE") and is_cover_page:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_logo = p_logo.add_run(placeholders["[Insérer Logo EMSI ici]"])
            r_logo.bold = True
            r_logo.font.color.rgb = RGBColor(255, 0, 0)
            
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_space = p2.add_run("\n\n")
            
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p3.add_run("RAPPORT DE PROJET DE FIN D'ANNÉE")
            run.bold = True
            run.font.size = Pt(24)
            continue
            
        if "PC Builder Maroc Plateforme de comparaison" in text or "PC Builder Maroc" in text and is_cover_page:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("PC Builder Maroc\nPlateforme de comparaison de prix et de vérification de compatibilité\n")
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.add_run("\n\n\n")
            continue
            
        if is_cover_page:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("École : EMSI Orangers, Casablanca\n").font.size = Pt(14)
            p.add_run("Filière : 3IIR — Ingénierie Informatique et Réseaux\n\n").font.size = Pt(14)
            p.add_run("Réalisé par : Salmane ELHJOUJI et Ghali KHARMOUDY\n").font.size = Pt(14)
            p.add_run("Encadrante : Prof. Houda Mouttalib\n\n").font.size = Pt(14)
            p.add_run("Année universitaire : 2025/2026\n").font.size = Pt(14)
            p.add_run("Date de soutenance : 14 Mai 2026\n").font.size = Pt(14)
            continue
            
        if text == "(Insérer Figure 3 : Planning ici)":
            p = doc.add_paragraph("Figure 3 : Planning du projet (Gantt/Tableau)", style='Heading 2')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = doc.add_table(rows=5, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Phase'
            hdr_cells[1].text = 'Tâches'
            hdr_cells[2].text = 'Responsable'
            hdr_cells[3].text = 'Durée'
            
            data = [
                ("Analyse", "UML, Spécifications", "Binôme", "2 semaines"),
                ("Back-end", "API, Scraping, DB", "Salmane", "4 semaines"),
                ("Front-end", "UI React, Configurateur", "Ghali", "4 semaines"),
                ("Validation", "Tests, Déploiement", "Binôme", "2 semaines")
            ]
            for r_idx, row in enumerate(data):
                cells = table.rows[r_idx+1].cells
                cells[0].text = row[0]
                cells[1].text = row[1]
                cells[2].text = row[2]
                cells[3].text = row[3]
            continue
            
        # Add placeholders instead of images
        if text in placeholders:
            box = doc.add_paragraph()
            box.alignment = WD_ALIGN_PARAGRAPH.CENTER
            box.paragraph_format.space_before = Pt(24)
            box.paragraph_format.space_after = Pt(24)
            r = box.add_run(placeholders[text])
            r.bold = True
            r.font.color.rgb = RGBColor(255, 102, 0) # Orange so it's obvious
            
            # If it's Figure 9, also inject the additional placeholders
            if text == "(Insérer Figure 9 : UI Configurateur PC ici)":
                for additional in additional_placeholders:
                    b2 = doc.add_paragraph()
                    b2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    b2.paragraph_format.space_before = Pt(12)
                    b2.paragraph_format.space_after = Pt(24)
                    r2 = b2.add_run(additional)
                    r2.bold = True
                    r2.font.color.rgb = RGBColor(255, 102, 0)
            continue
                
        # Skip Annexe D since it's now in chapter 5
        if text == "Annexe D : Résumé des tests automatisés" or text in ["Total Tests : 608", "Tests Unitaires : 412 (98% pass)", "Tests d'Intégration : 154 (100% pass)", "Tests de Propriétés (Fast-check) : 42 (100% pass)", "Couverture globale : 87% de la logique métier."]:
            continue

        doc.add_paragraph(text)
            
    # Add source code to Annexes to ensure length constraints
    doc.add_page_break()
    doc.add_paragraph("Annexe F : Extraits de Code Source Majeurs", style='Heading 1')
    
    for c_file in code_files:
        try:
            with open(c_file, 'r', encoding='utf-8') as cf:
                content = cf.read()
            doc.add_paragraph(f"Fichier : {c_file.split('/')[-1]}", style='Heading 2')
            
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
        except Exception as e:
            pass
            
    doc.save(output_file)
    print(f"Saved as {output_file}")

if __name__ == "__main__":
    code_files = [
        "c:/Headquarters/Projects/PcBuilder/shared/compatibility-engine.ts",
        "c:/Headquarters/Projects/PcBuilder/shared/types.ts"
    ]
    create_styled_doc(sys.argv[1], code_files, sys.argv[2])
