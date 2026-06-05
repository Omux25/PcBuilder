from docx import Document

def check_page_count(path):
    # This is a rough estimate as python-docx doesn't calculate pages
    # But we can check number of paragraphs and breaks
    doc = Document(path)
    breaks = 0
    paras = len(doc.paragraphs)
    for p in doc.paragraphs:
        if 'lastRenderedPageBreak' in p._element.xml or 'w:br w:type="page"' in p._element.xml:
            breaks += 1
    print(f"Total Paragraphs: {paras}")
    print(f"Manual Page Breaks: {breaks}")

if __name__ == "__main__":
    check_page_count(r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx')
