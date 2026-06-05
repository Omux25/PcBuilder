import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
import sys
import re

def create_styled_doc(input_text_file, images_dir, logo_path, code_files):
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        
    # Styles
    styles = doc.styles
    
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(12)
    
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(10)
    
    # Read paragraphs
    with open(input_text_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    paragraphs_text = []
    for line in lines:
        if line.startswith("["):
            match = re.match(r'\[\d+\] .*?: (.*)', line)
            if match:
                text = match.group(1).strip()
                if text:
                    paragraphs_text.append(text)
    
    markers = {
        "(Insérer Figure 1 : Campus EMSI Orangers ici)": ("image1.jpg", "Figure 1 : Campus EMSI Orangers"),
        "(Insérer Figure 2 : Organigramme ici)": ("image2.png", "Figure 2 : Organigramme"),
        "(Insérer Figure 4 : Diagramme de cas d'utilisation ici)": ("image3.png", "Figure 4 : Diagramme de cas d'utilisation"),
        "(Insérer Figure 5 : Diagramme de classes métier ici)": ("image4.png", "Figure 5 : Diagramme de classes métier"),
        "(Insérer Figure 6 : Schéma Relationnel / ERD ici)": ("image5.png", "Figure 6 : Schéma Relationnel / ERD"),
        "(Insérer Figure 7 : Règles de compatibilité ici)": ("image6.png", "Figure 7 : Règles de compatibilité"),
        "(Insérer Figure 8 : Pipeline de Scraping ici)": ("image7.png", "Figure 8 : Pipeline de Scraping"),
        "(Insérer Figure 9 : UI Configurateur PC ici)": ("image8.png", "Figure 9 : Aperçu fonctionnel des interfaces et Configurateur PC"),
        "Annexe B : Extrait du modèle relationnel (Insérer l'image du modèle de base de données ici)": ("image5.png", "Annexe B : Modèle relationnel")
    }
    
    page_break_before = [
        "Remerciements", "Résumé", "Abstract", "Table des matières", 
        "Liste des figures", "Introduction Générale", 
        "Chapitre 1 : Contexte du projet", "Chapitre 2 : Présentation du projet",
        "Chapitre 3 : Analyse et Conception", "Chapitre 4 : Réalisation technique",
        "Chapitre 5 : Résultats, tests et validation", "Conclusion et perspectives",
        "Bibliographie", "Annexes"
    ]
    
    h2_prefixes = [
        "1.1", "1.2", "1.3", "1.4",
        "2.1", "2.2", "2.3", "2.4",
        "3.1", "3.2", "3.3", "3.4",
        "4.1", "4.2", "4.3", "4.4", "4.5", "4.6",
        "5.1", "5.2", "5.3", "5.4", "5.5", "5.6"
    ]
    
    is_cover_page = True
    for text in paragraphs_text:
        
        is_h1 = text in page_break_before
        is_h2 = any(text.startswith(prefix + " ") for prefix in h2_prefixes)
        
        if is_cover_page and text == "Remerciements":
            is_cover_page = False
            
        if is_h1 and not is_cover_page:
            doc.add_page_break()
            p = doc.add_paragraph(text, style='Heading 1')
            continue
            
        if is_h2:
            p = doc.add_paragraph(text, style='Heading 2')
            continue
            
        if text.startswith("RAPPORT DE PROJET DE FIN D'ANNÉE"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                r_logo = p.add_run()
                r_logo.add_picture(logo_path, width=Inches(3.0))
            except:
                pass
            
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_space = p2.add_run("\n\n")
            
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p3.add_run("RAPPORT DE PROJET DE FIN D'ANNÉE")
            run.bold = True
            run.font.size = Pt(24)
            continue
            
        if "PC Builder Maroc Plateforme de comparaison" in text or "PC Builder Maroc" in text and is_cover_page:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("PC Builder Maroc\nPlateforme de comparaison de prix et de vérification de compatibilité\n")
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 51, 102)
            
            p.add_run("\n\n\n\n")
            continue
            
        if is_cover_page:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("École : EMSI Orangers, Casablanca\n").font.size = Pt(14)
            p.add_run("Filière : 3IIR — Ingénierie Informatique et Réseaux\n\n").font.size = Pt(14)
            p.add_run("Réalisé par : Salmane ELHJOUJI et Ghali KHARMOUDY\n").font.size = Pt(14)
            p.add_run("Encadrante : Prof. Houda Mouttalib\n\n").font.size = Pt(14)
            p.add_run("Année universitaire : 2025/2026\n").font.size = Pt(14)
            p.add_run("Date de soutenance : 14 Mai 2026\n").font.size = Pt(14)
            continue
            
        if text == "(Insérer Figure 3 : Planning ici)":
            p = doc.add_paragraph("Figure 3 : Planning du projet", style='Heading 2')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = doc.add_table(rows=5, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Phase'
            hdr_cells[1].text = 'Tâches'
            hdr_cells[2].text = 'Responsable'
            hdr_cells[3].text = 'Durée'
            
            data = [
                ("Analyse", "UML", "Binôme", "2 sem."),
                ("Back-end", "API", "Salmane", "4 sem."),
                ("Front-end", "UI", "Ghali", "4 sem."),
                ("Validation", "Tests", "Binôme", "2 sem.")
            ]
            for i, row in enumerate(data):
                cells = table.rows[i+1].cells
                cells[0].text = row[0]
                cells[1].text = row[1]
                cells[2].text = row[2]
                cells[3].text = row[3]
            continue
            
        inserted_image = False
        for marker, (img_filename, caption) in markers.items():
            if text == marker:
                img_path = f"{images_dir}/{img_filename}"
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run()
                    r.add_picture(img_path, width=Inches(6.0))
                    
                    cap_p = doc.add_paragraph(caption)
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.runs[0].italic = True
                except Exception as e:
                    print(f"Error inserting {img_filename}: {e}")
                inserted_image = True
                break
                
        if not inserted_image:
            doc.add_paragraph(text)
            
    # Add source code to Annexes to ensure length constraints
    doc.add_page_break()
    doc.add_paragraph("Annexe F : Extraits de Code Source", style='Heading 1')
    
    for c_file in code_files:
        try:
            with open(c_file, 'r', encoding='utf-8') as cf:
                content = cf.read()
            doc.add_paragraph(f"Fichier : {c_file.split('/')[-1]}", style='Heading 2')
            
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
        except Exception as e:
            print(f"Failed to read code file {c_file}: {e}")
            
    # Finally save
    doc.save("c:/Users/Omux2/Downloads/PcBuilder_Rapport_Final_V5.docx")
    print("Saved as PcBuilder_Rapport_Final_V5.docx")

if __name__ == "__main__":
    code_files = [
        "c:/Headquarters/Projects/PcBuilder/shared/compatibility-engine.ts",
        "c:/Headquarters/Projects/PcBuilder/shared/types.ts"
    ]
    create_styled_doc(sys.argv[1], sys.argv[2], sys.argv[3], code_files)
