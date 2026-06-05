import os
import shutil
from docx import Document
from docx.shared import Inches

def insert_images(docx_path, assets_map):
    # 1. Create Backup
    backup_path = docx_path.replace('.docx', '_BACKUP.docx')
    shutil.copy2(docx_path, backup_path)
    print(f"Backup created: {backup_path}")
    
    doc = Document(docx_path)
    
    # Track which images were inserted
    inserted_count = 0
    
    # Iterate through paragraphs to find placeholders
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Check against the mapping
        for placeholder, img_path in assets_map.items():
            if placeholder in text and os.path.exists(img_path):
                print(f"Found placeholder: '{placeholder}' -> Inserting {os.path.basename(img_path)}")
                
                # Clear the placeholder text or the whole paragraph
                para.text = text.replace(placeholder, "")
                
                # Insert image in the next paragraph or same
                run = para.add_run()
                run.add_picture(img_path, width=Inches(6.0)) # Standard width
                
                inserted_count += 1
                break
                
    # Also handle the mermaid-style placeholders by searching for the start of blocks
    # (Since they might span multiple paragraphs in some extractions, but here we just check for the headers)
    
    save_path = docx_path.replace('.docx', '_Final.docx')
    doc.save(save_path)
    print(f"Saved final report to: {save_path}")
    return save_path, inserted_count

if __name__ == "__main__":
    docx_file = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport.docx'
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    assets_dir = r'c:\Users\Omux2\Downloads\REPORT_ASSETS'
    
    mapping = {
        # EMSI General
        "Insérer ici une image de la façade du campus EMSI Orangers": os.path.join(assets_dir, "emsi_orangers_facade_1778638717569.png"),
        "Insérer ici l'Organigramme officiel": os.path.join(assets_dir, "emsi_organigramme_images_ddg_1778629388425.png"),
        
        # UML
        "3.2 Modélisation UML : Cas d'utilisation": os.path.join(puml_dir, "use_case.png"),
        "3.3 Architecture et Diagramme de Classes": os.path.join(puml_dir, "class.png"),
        
        # Technical Figures
        "Figure 7 : Lecture visuelle du moteur de compatibilité": os.path.join(puml_dir, "sequence_compatibility.png"),
        "Figure 8 : Chaîne de collecte automatique des prix": os.path.join(puml_dir, "sequence_scraping.png"),
        
        # App Screenshots
        "Figure 9 : Aperçu fonctionnel des interfaces": os.path.join(assets_dir, "figure_7_configurator_1778642579040.png"),
        "Figure 10 : Synthèse visuelle des résultats": os.path.join(assets_dir, "figure_10_admin_dashboard_1778643047831.png"),
        
        # Annexes
        "Annexe : diagramme UML des acteurs": os.path.join(puml_dir, "use_case.png"),
        "Annexe : séquence entre planificateur": os.path.join(puml_dir, "sequence_scraping.png")
    }
    
    insert_images(docx_file, mapping)
