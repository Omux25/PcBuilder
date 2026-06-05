import docx
import sys

def audit_docx(path):
    doc = docx.Document(path)
    print("--- AUDIT ---")
    
    empty_paras = 0
    run_fonts = {}
    run_sizes = {}
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            empty_paras += 1
            
        if "Table des matières" in text or "Liste des figures" in text:
            print(f"[{i}] {p.style.name} - {text}")
            # print runs in this para and next para
            try:
                next_p = doc.paragraphs[i+1]
                print(f"  Next para text: '{next_p.text}'")
                print(f"  Next para XML: {next_p._p.xml[:200]}...")
            except:
                pass
                
        for r in p.runs:
            font_name = r.font.name if r.font.name else (p.style.font.name if p.style.font.name else "Default")
            font_size = r.font.size.pt if r.font.size else (p.style.font.size.pt if p.style.font.size else "Default")
            
            run_fonts[font_name] = run_fonts.get(font_name, 0) + 1
            run_sizes[font_size] = run_sizes.get(font_size, 0) + 1
            
    print(f"Empty paragraphs: {empty_paras}")
    print(f"Fonts used: {run_fonts}")
    print(f"Font sizes used: {run_sizes}")

if __name__ == "__main__":
    audit_docx(sys.argv[1])
