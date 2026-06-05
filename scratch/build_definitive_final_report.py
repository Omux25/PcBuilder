import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_definitive_final_report():
    downloads = r'c:\Users\Omux2\Downloads'
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    assets_dir = os.path.join(downloads, 'REPORT_ASSETS')
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    
    with open(r'c:\Headquarters\Projects\PcBuilder\scratch\final_text_audit.txt', 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    print("Building DEFINITIVE FINAL REPORT...")

    # Title Page
    def add_centered(text, size=12, bold=False, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size); run.bold = bold
        p.paragraph_format.space_after = Pt(space_after)

    add_centered("RAPPORT DE PROJET DE FIN D'ANNÉE", 24, True, 20)
    add_centered("PC Builder Maroc", 22, True, 10)
    add_centered("Plateforme de comparaison de prix et de vérification de compatibilité", 14, False, 40)
    add_centered("École : EMSI Orangers, Casablanca", 12, False, 5)
    add_centered("Filière : 3IIR — Ingénierie Informatique et Réseaux", 12, False, 20)
    add_centered("Étudiants : Salmane ELHJOUJI · Ghali KHARMOUDY", 12, False, 5)
    add_centered("Encadrante : Prof. Houda Mouttalib", 12, False, 20)
    add_centered("Année : 2025/2026", 12)
    doc.add_page_break()

    # MANUAL MAPPING BASED ON AUDIT
    sections = [
        ("Remerciements", lines[10:16]),
        ("Résumé", lines[17:21]),
        ("Abstract", lines[22:24]),
        ("Table des matières", []),
        ("Introduction Générale", lines[32:36]), # Line 37 is a duplicate, SKIPPED.
        ("Chapitre 1 : Contexte du projet", []),
        ("1.1 Présentation de l’EMSI", lines[41:46]),
        ("1.2 Structure Organisationnelle", lines[48:51]),
        ("1.3 Contexte du marché marocain", lines[53:56]),
        ("1.4 Problèmes observés", lines[56:61]),
        ("Chapitre 2 : Présentation du projet", []),
        ("2.1 Problématique", lines[63:67]),
        ("2.2 Cahier des Charges", lines[67:85]),
        ("2.3 Solution Proposée", lines[85:87]),
        ("2.4 Planning", []), # Line 87 is the header, line 88 is Fig caption, line 90 is STRAY header.
        ("Chapitre 3 : Analyse et Conception", []),
        ("3.1 Justification des Choix Techniques", lines[91:96]),
        ("3.2 Modélisation UML : Cas d'utilisation", []),
        ("3.3 Architecture et Diagramme de Classes", []),
        ("3.4 Conception de la Base de Données", []),
        ("Chapitre 4 : Réalisation technique", []),
        ("4.1 Environnement technique", lines[107:113]),
        ("4.2 Serveur applicatif et API REST", lines[113:119]),
        ("4.3 Moteur de compatibilité", lines[119:126]),
        ("4.4 Système de collecte automatique", lines[128:135]),
        ("4.5 Interface utilisateur", lines[137:143]), # Line 144 is duplicate Figure 9, SKIPPED.
        ("4.6 Panneau d'administration", lines[145:151]),
        ("Chapitre 5 : Résultats, tests et validation", []),
        ("5.1 Fonctionnalités livrées", lines[153:159]),
        ("5.2 Tests réalisés", lines[159:165]),
        ("5.3 Apport de la solution", lines[165:171]),
        ("5.4 Limites identifiées", lines[171:175]),
        ("5.5 Organisation du travail", lines[175:179]),
        ("5.6 Perspectives techniques détaillées", lines[179:183]),
        ("Conclusion et perspectives", lines[184:191]),
        ("Bibliographie", lines[192:206]),
        ("Annexes", lines[207:228])
    ]

    fig_count = 1
    def add_fig(path, title, width=5.0):
        nonlocal fig_count
        if os.path.exists(path):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(width))
            cap = doc.add_paragraph(f"Figure {fig_count} : {title}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_count += 1

    for title, content in sections:
        if "Chapitre" in title or title in ["Conclusion et perspectives", "Bibliographie", "Annexes"]:
            doc.add_page_break()
            doc.add_heading(title, level=1)
        elif title == "Table des matières":
            doc.add_heading(title, level=1)
            doc.add_paragraph("Introduction Générale ................................................. 1")
            doc.add_paragraph("Chapitre 1 : Contexte ................................................ 3")
            doc.add_paragraph("Chapitre 2 : Présentation du projet ................................ 9")
            doc.add_paragraph("Chapitre 3 : Analyse et Conception ................................ 16")
            doc.add_paragraph("Chapitre 4 : Réalisation Technique ................................. 24")
            doc.add_paragraph("Chapitre 5 : Tests et Validation .................................... 32")
        else:
            doc.add_heading(title, level=2)
        
        for line in content:
            txt = line.strip()
            # MANUALLY FILTER STRAY JUNK
            if re.match(r'^\d+\.\d+\s', txt): continue
            if txt.lower() == title.lower(): continue
            if txt.startswith("Figure "): continue
            if "Chapitre" in txt and len(txt) < 40: continue # Catch stray headers
            if txt: doc.add_paragraph(txt)
            
        if "EMSI" in title: add_fig(os.path.join(downloads, "EMSI-Maroc-0x280.jpg"), "Campus EMSI Orangers")
        if "Organisationnelle" in title: add_fig(os.path.join(puml_dir, "emsi_organigramme.png"), "Organigramme Institutionnel")
        if "Planning" in title: 
            table = doc.add_table(rows=5, cols=4); table.style = 'Table Grid'
            data = [["Phase", "Tâches", "Responsable", "Durée"], ["Analyse", "UML", "Binôme", "2 sem."], ["Back-end", "API", "Salmane", "4 sem."], ["Front-end", "UI", "Ghali", "4 sem."], ["Validation", "Tests", "Binôme", "2 sem."]]
            for r, row in enumerate(data):
                for c, val in enumerate(row): table.cell(r, c).text = val
            doc.add_paragraph(f"Figure {fig_count} : Planning GANTT").alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_count += 1
        if "Cas d'utilisation" in title: add_fig(os.path.join(puml_dir, "use_case.png"), "Diagramme de cas d'utilisation")
        if "Architecture" in title and "Classes" in title: add_fig(os.path.join(puml_dir, "class.png"), "Diagramme de classes métier")
        if "Base de Données" in title: add_fig(os.path.join(puml_dir, "database_erd.png"), "Schéma Relationnel de la Base de Données")
        if "collecte automatique" in title: add_fig(os.path.join(puml_dir, "sequence_scraping.png"), "Pipeline de Scraping")
        if "compatibilité" in title: add_fig(os.path.join(puml_dir, "sequence_compatibility.png"), "Vérification des règles de compatibilité")
        if "Interface utilisateur" in title: add_fig(os.path.join(assets_dir, "figure_7_home_1778642725154.png"), "Configurateur PC - Interface Principale")

    for section in doc.sections:
        section.top_margin = Inches(0.98); section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98); section.right_margin = Inches(0.98)
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1.5
        if len(p.text) > 80: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            if not run.bold: run.font.size = Pt(12)

    doc.save(final_path)
    print("Definitive Final Report Built.")

if __name__ == "__main__":
    build_definitive_final_report()
