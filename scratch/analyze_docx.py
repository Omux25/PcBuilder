import docx
import sys
import os

def analyze_docx(path):
    print(f"--- Analyzing {path} ---")
    doc = docx.Document(path)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Inline shapes (Images): {len(doc.inline_shapes)}")
    
    # Check for images in paragraphs
    img_count = 0
    for p in doc.paragraphs:
        if 'Graphic' in p._p.xml or 'pic:pic' in p._p.xml:
            img_count += 1
    print(f"Images in XML: {img_count}")
    
    # Print first few paragraphs to see
    print("First 10 non-empty paragraphs:")
    count = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            print(f"  {text[:100]}")
            count += 1
            if count == 10:
                break
    print("\n")

if __name__ == "__main__":
    analyze_docx(sys.argv[1])
    analyze_docx(sys.argv[2])
