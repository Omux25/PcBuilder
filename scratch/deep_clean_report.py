import os
import shutil
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_and_insert(docx_path, assets_map):
    # 1. Create a fresh Backup from the original
    original_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport.docx'
    final_path = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    shutil.copy2(original_path, final_path)
    
    doc = Document(final_path)
    
    print(f"Starting deep clean and insertion on {final_path}")
    
    # We need to iterate and potentially delete paragraphs
    # It's easier to track indices and delete from end to start or just mark for deletion
    
    paragraphs = doc.paragraphs
    to_delete = []
    
    in_mermaid = False
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        
        # Detect Mermaid blocks
        if text.startswith("```mermaid") or text == "```":
            in_mermaid = not in_mermaid
            to_delete.append(i)
            continue
        
        if in_mermaid:
            to_delete.append(i)
            continue
            
        # Check for placeholders and instructions
        for placeholder, img_path in assets_map.items():
            if placeholder in para.text:
                print(f"Matched: {placeholder}")
                # Clear instruction text
                para.text = para.text.replace(placeholder, "")
                
                # If image exists, insert it
                if os.path.exists(img_path):
                    run = para.add_run()
                    run.add_break()
                    run.add_picture(img_path, width=Inches(5.5))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"Inserted: {os.path.basename(img_path)}")
                else:
                    print(f"Warning: Image not found at {img_path}")
                break

    # Delete marked paragraphs (in reverse order to maintain indices)
    for index in sorted(to_delete, reverse=True):
        p = doc.paragraphs[index]
        parent = p._element.getparent()
        parent.remove(p._element)

    doc.save(final_path)
    print(f"Deep clean complete. Saved to {final_path}")

if __name__ == "__main__":
    assets_dir = r'c:\Users\Omux2\Downloads\REPORT_ASSETS'
    puml_dir = r'C:\Headquarters\Studies\notes\diagrams\rendered'
    
    assets_map = {
        # EMSI General
        "Insérer ici une image de la façade du campus EMSI Orangers": os.path.join(assets_dir, "emsi_orangers_facade_1778638717569.png"),
        "Insérer ici l'Organigramme officiel": os.path.join(assets_dir, "emsi_organigramme_images_ddg_1778629388425.png"),
        
        # UML Sections (The images will be placed right after these headings)
        "3.2 Modélisation UML : Cas d'utilisation": os.path.join(puml_dir, "use_case.png"),
        "3.3 Architecture et Diagramme de Classes": os.path.join(puml_dir, "class.png"),
        "3.4 Conception de la Base de Données": os.path.join(assets_dir, "figure_6_database_erd_1778640799727.png"),
        
        # Technical Figures
        "Figure 7 : Lecture visuelle du moteur de compatibilité": os.path.join(puml_dir, "sequence_compatibility.png"),
        "Figure 8 : Chaîne de collecte automatique des prix": os.path.join(puml_dir, "sequence_scraping.png"),
        
        # App Screenshots
        "Figure 9 : Aperçu fonctionnel des interfaces": os.path.join(assets_dir, "figure_7_home_1778642725154.png"),
        "Figure 10 : Synthèse visuelle des résultats": os.path.join(assets_dir, "figure_10_admin_dashboard_1778643047831.png"),
        
        # Annexes
        "Annexe : diagramme UML des acteurs": os.path.join(puml_dir, "use_case.png"),
        "Annexe : séquence entre planificateur": os.path.join(puml_dir, "sequence_scraping.png")
    }
    
    clean_and_insert(r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx', assets_map)
