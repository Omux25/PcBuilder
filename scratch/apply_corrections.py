import docx
import sys
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_teacher_corrections(input_path, output_path):
    doc = docx.Document(input_path)

    # 1. Replace teacher name
    for p in doc.paragraphs:
        if "Houda Mouttalib" in p.text:
            p.text = p.text.replace("Prof. Houda Mouttalib", "Prof. Soukayna MAYARA")
            p.text = p.text.replace("Houda Mouttalib", "Soukayna MAYARA")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Fix misstyled headings (TOC too long)
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') or p.style.name.startswith('Titre'):
            if len(p.text.strip()) > 120:
                print(f"Fixing long heading: {p.text[:50]}...")
                p.style = 'Normal'

    # 3. Add Global Architecture Diagram Placeholder
    inserted_arch = False
    for p in doc.paragraphs:
        if p.text.startswith("3.1"): 
            if not inserted_arch:
                arch_title = p.insert_paragraph_before("3.X Schéma d'Architecture Global", style='Heading 2')
                
                arch_body = p.insert_paragraph_before("Cette section présente l'architecture technique globale du projet, illustrant les flux de données entre le frontend (React), le backend (Hono/Bun), la base de données (PostgreSQL) et le moteur de scraping.")
                arch_body.style = 'Normal'
                
                arch_box = p.insert_paragraph_before()
                r = arch_box.add_run("[PLACEHOLDER IMAGE: Insérer ici un Schéma d'Architecture Global détaillant les interactions techniques du système]")
                r.bold = True
                r.font.color.rgb = RGBColor(255, 102, 0)
                arch_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                inserted_arch = True
                break

    # 4. Add limits of scraping
    # Let's append scraping limits right before Chapter 5
    inserted_scraping = False
    for p in doc.paragraphs:
        if "Chapitre 5" in p.text and (p.style.name.startswith('Heading 1') or p.style.name.startswith('Titre 1')):
            if not inserted_scraping:
                scrap_title = p.insert_paragraph_before("4.X Limites juridiques et techniques du scraping", style='Heading 2')
                
                scrap_body = p.insert_paragraph_before(
                    "Il est crucial de souligner les limites inhérentes à la collecte automatisée de données. "
                    "Sur le plan technique, notre pipeline doit faire face aux changements imprévisibles de la structure HTML des sites sources, "
                    "aux erreurs réseau, ainsi qu'aux mécanismes de protection anti-bot (tels que les blocages d'IP ou les CAPTCHAs). "
                    "Afin de minimiser ces risques, la fréquence de collecte est régulée pour ne pas surcharger les serveurs cibles. "
                    "Sur le plan juridique et éthique, le projet s'inscrit dans une démarche de respect des conditions d'utilisation "
                    "des revendeurs, en limitant l'extraction aux données publiques de tarification, tout en veillant à ne pas causer de "
                    "préjudice aux plateformes source."
                )
                scrap_body.style = 'Normal'
                inserted_scraping = True
                break

    # 5. Remove placeholder in Annex B
    for p in doc.paragraphs:
        if "Capture d'écran détaillée du Schéma de Base de données" in p.text:
            p.text = ""

    doc.save(output_path)
    print(f"Modifications saved to {output_path}")

if __name__ == "__main__":
    apply_teacher_corrections(sys.argv[1], sys.argv[2])
