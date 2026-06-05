import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def test_toc():
    doc = docx.Document()
    doc.add_heading("Table des matières", 1)
    add_toc(doc)
    
    doc.add_page_break()
    doc.add_heading("Introduction", 1)
    doc.add_paragraph("Hello world")
    
    doc.save("scratch/test_toc.docx")

if __name__ == "__main__":
    test_toc()
