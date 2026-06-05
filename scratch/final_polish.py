from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def spacing_and_polish_fix(docx_path):
    doc = Document(docx_path)
    
    print(f"Applying final spacing polish to {docx_path}")
    
    # 1. First, identify and remove "Unknown diagram error" paragraphs
    to_delete = []
    for i, para in enumerate(doc.paragraphs):
        if "Unknown diagram error" in para.text or "[INSTRUCTION" in para.text:
            to_delete.append(i)
            
    # Delete from bottom to top
    for index in sorted(to_delete, reverse=True):
        p = doc.paragraphs[index]
        parent = p._element.getparent()
        parent.remove(p._element)

    # 2. Iterate through all remaining paragraphs to fix alignment/spacing
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # If it's a heading or part of Annexes, use LEFT alignment
        is_annexe = "Annexe" in text or text.startswith("1.") or text.startswith("2.") or text.startswith("3.")
        is_list_item = text.startswith("-") or text.startswith("•") or (len(text) < 100 and ":" in text)
        
        if is_annexe or is_list_item:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Main body text: keep Justified but strip internal \n to prevent stretching
            if len(text) > 150:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Remove manual line breaks within the paragraph text
                para.text = text.replace("\n", " ").replace("  ", " ")
            else:
                # Short lines should NOT be justified
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Global font enforcement (Times New Roman 12)
        for run in para.runs:
            run.font.name = 'Times New Roman'
            if not run.bold:
                run.font.size = Pt(12)

    # 3. Fix Table formatting (Gantt)
    for table in doc.tables:
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)

    doc.save(docx_path)
    print("Polish complete. Spacing fixed and error artifacts removed.")

if __name__ == "__main__":
    final_docx = r'c:\Users\Omux2\Downloads\PcBuilder_Rapport_Final.docx'
    spacing_and_polish_fix(final_docx)
